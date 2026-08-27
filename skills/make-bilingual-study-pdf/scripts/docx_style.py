#!/usr/bin/env python3
from __future__ import annotations

import argparse
import io
import os
import re
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

from profile import profile_contract
from safe_artifacts import (
    ArtifactSafetyError,
    atomic_publish_with_writer,
    lexical_absolute_path,
    prepare_artifact_directory,
    read_artifact_bytes,
    validate_artifact_directory,
    validate_artifact_file,
    validate_artifact_tree,
    work_relative_artifact_path,
)
from semantic_registry import get_style


LATIN_FONT = "Noto Sans"
CJK_FONT = "Noto Sans S Chinese"
CODE_FONT = "DejaVu Sans Mono"
INK = "172033"
MUTED = "566574"
ACCENT = "2F6073"
ACCENT_LIGHT = "D9E7EC"
ZH_FILL = "F3F7F9"
PROBLEM = "D97706"
PROBLEM_FILL = "FFFBF5"
TIP = "4D7C8A"
TIP_FILL = "F3F8FA"
EXAMPLE = "708890"
EXAMPLE_FILL = "F6F8F9"
TARGET_TEXT_RE = re.compile(r"[\u3400-\u9fff]")
CALLOUT_SOURCE_PATTERNS = {
    "problem": re.compile(r"^Problem\s*[（(]", re.I),
    "example": re.compile(r"^Example\s*[（(]", re.I),
    "tip": re.compile(r"^Low-Resource Tip(?:\s*[:：]|$)", re.I),
}
CALLOUT_TARGET_PATTERNS = {
    "problem": re.compile(r"^问题\s*[（(]", re.I),
    "example": re.compile(r"^示例\s*[（(]", re.I),
    "tip": re.compile(r"^低资源提示(?:\s*[:：]|$)", re.I),
}
PROBLEM_BEGIN = "V2-PROBLEM-CALLOUT-BEGIN"
PROBLEM_END = "V2-PROBLEM-CALLOUT-END"
GENERIC_BEGIN = "V23-CALLOUT-BEGIN"
GENERIC_END = "V23-CALLOUT-END"
ACTIVE_SCHEMA_VERSION = 1
ROLE_STYLES: dict[str, str] = {}
CALLOUT_RANGE_IDS: dict[tuple[int, int, str], str] = {}
CALLOUT_RANGE_MEMBERSHIPS: dict[tuple[int, int, str], str] = {}
GENERIC_PALETTES = {
    "abstract": ("4D7C8A", "F3F8FA"),
    "definition": ("2F7D5B", "F3FAF6"),
    "theorem": ("365E9D", "F4F7FC"),
    "proof": ("667085", "F7F8FA"),
    "note": ("4D7C8A", "F3F8FA"),
    "warning": ("B54708", "FFF8F0"),
    "exercise": ("7A5AF8", "F7F5FF"),
}


def _bounded_cli_path(boundary: Path, path: Path, *, label: str) -> Path:
    boundary = lexical_absolute_path(boundary)
    candidate = lexical_absolute_path(path)
    try:
        relative = os.path.relpath(candidate, boundary)
    except ValueError as exc:
        raise ArtifactSafetyError(f"{label} is outside WORK") from exc
    bounded = work_relative_artifact_path(
        boundary, Path(relative).as_posix(), label=label
    )
    if os.path.normcase(os.fspath(bounded)) != os.path.normcase(os.fspath(candidate)):
        raise ArtifactSafetyError(f"{label} is outside WORK")
    return bounded


def configure_profile(profile: dict) -> None:
    global TARGET_TEXT_RE, CALLOUT_SOURCE_PATTERNS, CALLOUT_TARGET_PATTERNS
    global ACTIVE_SCHEMA_VERSION, ROLE_STYLES
    TARGET_TEXT_RE = re.compile(profile["translation"]["target_text_pattern"])
    ACTIVE_SCHEMA_VERSION = int(profile.get("schema_version", 1))
    if ACTIVE_SCHEMA_VERSION == 1:
        CALLOUT_SOURCE_PATTERNS = {
            group["style"]: re.compile(group["source_pattern"], re.I)
            for group in profile["semantics"]["groups"]
        }
        CALLOUT_TARGET_PATTERNS = {
            group["style"]: re.compile(group["target_pattern"], re.I)
            for group in profile["semantics"]["groups"]
        }
        ROLE_STYLES = {group["role"]: group["style"] for group in profile["semantics"]["groups"]}
        return
    contract = profile_contract(profile)
    ROLE_STYLES = {item["role"]: item["style"] for item in contract["roles"]}
    CALLOUT_SOURCE_PATTERNS = {}
    CALLOUT_TARGET_PATTERNS = {}
    for item in contract["roles"]:
        if item["grouping"] != "structural-container":
            continue
        for selector in item["selectors"]:
            if "source_pattern" in selector:
                CALLOUT_SOURCE_PATTERNS.setdefault(
                    item["role"], re.compile(selector["source_pattern"], re.I)
                )
            if "target_pattern" in selector:
                CALLOUT_TARGET_PATTERNS.setdefault(
                    item["role"], re.compile(selector["target_pattern"], re.I)
                )


def has_cjk(text: str) -> bool:
    return bool(TARGET_TEXT_RE.search(text))


def callout_role(text: str, *, include_target: bool = True) -> str | None:
    stripped = text.strip()
    for role, pattern in CALLOUT_SOURCE_PATTERNS.items():
        if pattern.search(stripped):
            return role
    if include_target:
        for role, pattern in CALLOUT_TARGET_PATTERNS.items():
            if pattern.search(stripped):
                return role
    return None


def set_run_font(run, name: str, size: float, *, color: str = INK, bold: bool | None = None) -> None:
    run.font.name = name
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.get_or_add_rFonts()
    r_fonts.set(qn("w:ascii"), name)
    r_fonts.set(qn("w:hAnsi"), name)
    r_fonts.set(qn("w:eastAsia"), CJK_FONT if name != CODE_FONT else CODE_FONT)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold


def set_spacing(paragraph, *, before: float = 0, after: float = 0, line: float = 1.2) -> None:
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def clear_shading_and_borders(paragraph) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    for tag in ("w:shd", "w:pBdr"):
        node = p_pr.find(qn(tag))
        if node is not None:
            p_pr.remove(node)


def add_shading(paragraph, fill: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def add_paragraph_border(
    paragraph,
    *,
    color: str,
    left: bool = False,
    right: bool = False,
    top: bool = False,
    bottom: bool = False,
    size: int = 10,
    space: int = 7,
) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    borders = p_pr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        p_pr.append(borders)
    for edge, enabled in (("left", left), ("right", right), ("top", top), ("bottom", bottom)):
        if not enabled:
            continue
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), str(size))
        node.set(qn("w:space"), str(space))
        node.set(qn("w:color"), color)


def clear_paragraph_content(paragraph) -> None:
    """Remove rendered runs/drawings while preserving paragraph properties."""
    for child in list(paragraph._p):
        if child.tag != qn("w:pPr"):
            paragraph._p.remove(child)


def is_numbered_paragraph(paragraph) -> bool:
    p_pr = paragraph._p.pPr
    return p_pr is not None and p_pr.find(qn("w:numPr")) is not None


def has_legacy_horizontal_rule(paragraph) -> bool:
    return bool(
        paragraph._p.xpath(
            ".//*[local-name()='rect' and @*[local-name()='hr']='t']"
        )
    )


def is_heading(paragraph) -> bool:
    return paragraph.style.name.startswith("Heading ")


def is_bold_label(paragraph) -> bool:
    nonempty = [run for run in paragraph.runs if run.text.strip()]
    return bool(nonempty) and all(run.bold for run in nonempty) and len(paragraph.text.strip()) <= 90


def set_keep_with_next(paragraph, value: bool) -> None:
    paragraph.paragraph_format.keep_with_next = value


def style_heading(paragraph) -> None:
    level = int(paragraph.style.name.rsplit(" ", 1)[-1])
    sizes = {1: 21.0, 2: 15.0, 3: 13.0, 4: 11.5}
    before = {1: 0.0, 2: 17.0, 3: 13.0, 4: 10.0}
    after = {1: 7.0, 2: 5.0, 3: 4.0, 4: 3.0}
    size = sizes.get(level, 11.0)
    set_spacing(paragraph, before=before.get(level, 8), after=after.get(level, 3), line=1.08)
    set_keep_with_next(paragraph, True)
    for run in paragraph.runs:
        set_run_font(run, LATIN_FONT, size, color=ACCENT if level != 2 else INK, bold=True)


def style_heading_translation(paragraph, level: int | None) -> None:
    sizes = {1: 13.5, 2: 11.3, 3: 10.8, 4: 10.4}
    set_spacing(paragraph, after=7 if level != 1 else 12, line=1.15)
    set_keep_with_next(paragraph, True)
    paragraph.paragraph_format.left_indent = Inches(0)
    paragraph.paragraph_format.right_indent = Inches(0)
    clear_shading_and_borders(paragraph)
    for run in paragraph.runs:
        set_run_font(run, CJK_FONT, sizes.get(level or 4, 10.4), color=ACCENT, bold=True)


def style_english_body(paragraph) -> None:
    set_spacing(paragraph, after=3, line=1.2)
    clear_shading_and_borders(paragraph)
    for run in paragraph.runs:
        if run.style and run.style.name == "Verbatim Char":
            set_run_font(run, CODE_FONT, 9.4, color=INK)
        else:
            set_run_font(run, LATIN_FONT, 10.5, color=INK)


def style_chinese_body(paragraph) -> None:
    set_spacing(paragraph, after=9, line=1.3)
    if paragraph.style.name not in {"Compact", "Normal"}:
        paragraph.paragraph_format.left_indent = Inches(0.18)
        paragraph.paragraph_format.right_indent = Inches(0.06)
    clear_shading_and_borders(paragraph)
    add_shading(paragraph, ZH_FILL)
    add_paragraph_border(paragraph, color=ACCENT_LIGHT, left=True, size=10, space=7)
    for run in paragraph.runs:
        if run.style and run.style.name == "Verbatim Char":
            set_run_font(run, CODE_FONT, 9.2, color=INK)
        else:
            set_run_font(run, CJK_FONT, 10.2, color="334653")


def style_code(paragraph) -> None:
    set_spacing(paragraph, before=4, after=8, line=1.08)
    paragraph.paragraph_format.left_indent = Inches(0.16)
    paragraph.paragraph_format.right_indent = Inches(0.06)
    clear_shading_and_borders(paragraph)
    add_shading(paragraph, "F4F6F8")
    add_paragraph_border(paragraph, color="C6D1D6", left=True, size=10, space=7)
    for run in paragraph.runs:
        set_run_font(run, CODE_FONT, 8.4, color=INK)


def style_math(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_spacing(paragraph, before=4, after=8, line=1.0)


def style_table(table) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    table_properties = table._tbl.tblPr
    borders = table_properties.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        table_properties.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = borders.find(qn(f"w:{edge}"))
        if border is None:
            border = OxmlElement(f"w:{edge}")
            borders.append(border)
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "6")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), "9FB3BD")
    for row_index, row in enumerate(table.rows):
        if row_index == 0:
            tr_pr = row._tr.get_or_add_trPr()
            repeat = OxmlElement("w:tblHeader")
            repeat.set(qn("w:val"), "true")
            tr_pr.append(repeat)
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if row_index == 0:
                tc_pr = cell._tc.get_or_add_tcPr()
                shading = tc_pr.find(qn("w:shd"))
                if shading is None:
                    shading = OxmlElement("w:shd")
                    tc_pr.append(shading)
                shading.set(qn("w:fill"), ACCENT_LIGHT)
            for paragraph in cell.paragraphs:
                set_spacing(paragraph, before=0, after=2, line=1.08)
                for run in paragraph.runs:
                    set_run_font(
                        run,
                        LATIN_FONT,
                        9.2,
                        color=INK,
                        bold=True if row_index == 0 else None,
                    )


def is_math_paragraph(paragraph) -> bool:
    return bool(paragraph._p.xpath(".//m:oMath | .//m:oMathPara"))


def find_callout_ranges(paragraphs) -> tuple[list[tuple[int, int, str]], list[int]]:
    global CALLOUT_RANGE_IDS, CALLOUT_RANGE_MEMBERSHIPS
    CALLOUT_RANGE_IDS = {}
    CALLOUT_RANGE_MEMBERSHIPS = {}
    if ACTIVE_SCHEMA_VERSION == 2:
        ranges: list[tuple[int, int, str]] = []
        marker_indices: list[int] = []
        active: tuple[int, str, str, str, str] | None = None
        for marker_index, paragraph in enumerate(paragraphs):
            text = paragraph.text.strip()
            if text.startswith(f"{GENERIC_BEGIN}|"):
                parts = text.split("|", 4)
                if len(parts) != 5 or active is not None:
                    raise ValueError("invalid or nested generic callout begin marker")
                _, role, style, membership, group_id = parts
                if membership not in {"complete", "anchor-only"}:
                    raise ValueError(f"invalid generic callout membership: {membership}")
                if ROLE_STYLES.get(role) != style:
                    raise ValueError(f"generic callout marker style mismatch for role {role}")
                style_spec = get_style(style)
                if not style_spec.supports_structural_container:
                    raise ValueError(f"role {role} style cannot form a structural container")
                active = (marker_index, role, style, membership, group_id)
                marker_indices.append(marker_index)
                continue
            if text.startswith(f"{GENERIC_END}|"):
                parts = text.split("|", 4)
                if len(parts) != 5 or active is None:
                    raise ValueError("unmatched generic callout end marker")
                start_marker, role, style, membership, group_id = active
                if parts[1:] != [role, style, membership, group_id]:
                    raise ValueError(f"generic callout marker pair mismatch: {group_id}")
                if marker_index <= start_marker + 1:
                    raise ValueError(f"empty generic callout range: {group_id}")
                item = (start_marker + 1, marker_index - 1, role)
                ranges.append(item)
                CALLOUT_RANGE_IDS[item] = group_id
                CALLOUT_RANGE_MEMBERSHIPS[item] = membership
                marker_indices.append(marker_index)
                active = None
        if active is not None:
            raise ValueError(f"unclosed generic callout marker: {active[4]}")
        return ranges, marker_indices

    def starts_callout(paragraph) -> bool:
        # Target-language labels belong to the current bilingual callout. Only
        # source-language labels or explicit range markers may open a new range.
        return callout_role(paragraph.text, include_target=False) is not None

    ranges: list[tuple[int, int, str]] = []
    marker_indices: list[int] = []
    problem_ranges: list[tuple[int, int, str]] = []
    active_problem: int | None = None
    for marker_index, marker_paragraph in enumerate(paragraphs):
        marker_text = marker_paragraph.text.strip()
        if marker_text == PROBLEM_BEGIN:
            active_problem = marker_index
            marker_indices.append(marker_index)
        elif marker_text == PROBLEM_END and active_problem is not None:
            marker_indices.append(marker_index)
            if marker_index > active_problem + 1:
                problem_ranges.append((active_problem + 1, marker_index - 1, "problem"))
            active_problem = None
    ranges.extend(problem_ranges)
    problem_indices = {
        item_index for start, end, _ in problem_ranges for item_index in range(start, end + 1)
    }
    index = 0
    while index < len(paragraphs):
        paragraph = paragraphs[index]
        text = paragraph.text.strip()
        if index in marker_indices or index in problem_indices:
            index += 1
            continue
        if paragraph.style.name != "Block Text":
            index += 1
            continue
        role = callout_role(text, include_target=False)
        if role == "problem" and problem_ranges:
            role = None
        if role is None:
            index += 1
            continue
        end = index
        while end + 1 < len(paragraphs):
            candidate = paragraphs[end + 1]
            if candidate.style.name == "Block Text":
                if starts_callout(candidate):
                    break
                end += 1
                continue
            if (
                not candidate.text.strip()
                and end + 2 < len(paragraphs)
                and paragraphs[end + 2].style.name == "Block Text"
                and not starts_callout(paragraphs[end + 2])
            ):
                end += 1
                continue
            break
        ranges.append((index, end, role))
        index = end + 1
    ranges.sort(key=lambda item: item[0])
    return ranges, marker_indices


def style_callout(paragraphs, start: int, end: int, role: str) -> None:
    style_id = ROLE_STYLES.get(role, role)
    color, fill = {
        "problem": (PROBLEM, PROBLEM_FILL),
        "tip": (TIP, TIP_FILL),
        "example": (EXAMPLE, EXAMPLE_FILL),
        **GENERIC_PALETTES,
    }[style_id]
    nonempty = [index for index in range(start, end + 1) if paragraphs[index].text.strip()]
    first = nonempty[0]
    last = nonempty[-1]
    seen_chinese = False
    separator_index = next((index for index in range(start, end + 1) if not paragraphs[index].text.strip()), None)
    for index in range(start, end + 1):
        paragraph = paragraphs[index]
        text = paragraph.text.strip()
        clear_shading_and_borders(paragraph)
        add_shading(paragraph, fill)
        paragraph.paragraph_format.left_indent = Inches(0.12)
        paragraph.paragraph_format.right_indent = Inches(0.12)
        if role == "problem" and is_numbered_paragraph(paragraph):
            # Pandoc numbering contributes a hanging indent even after a direct
            # left indent is set. Paragraph borders follow that inherited hanging
            # origin, producing a stepped outline. Explicit zero first-line indent
            # keeps the real numbering while aligning every border fragment.
            paragraph.paragraph_format.first_line_indent = Inches(0)
        if not text:
            set_spacing(paragraph, before=2, after=3, line=1.0)
            if role == "problem" or ACTIVE_SCHEMA_VERSION == 2:
                clear_paragraph_content(paragraph)
            add_paragraph_border(
                paragraph, color=color, left=True, right=True, top=True, size=8, space=6
            )
            continue
        cjk = has_cjk(text)
        if cjk:
            seen_chinese = True
        set_spacing(paragraph, before=4 if index == first else 0, after=5 if index == last else 2, line=1.22)
        add_paragraph_border(
            paragraph,
            color=color,
            left=True,
            right=True,
            top=index == first,
            bottom=index == last,
            size=12 if role == "problem" else 10,
            space=8,
        )
        heading = index == first or (
            (role == "problem" or ACTIVE_SCHEMA_VERSION == 2)
            and cjk
            and (separator_index is not None and index == separator_index + 1)
        )
        for run in paragraph.runs:
            if run.style and run.style.name == "Verbatim Char":
                set_run_font(run, CODE_FONT, 9.1, color=INK)
            else:
                set_run_font(run, CJK_FONT if cjk else LATIN_FONT, 10.0, color=ACCENT if cjk else INK, bold=True if heading else None)
        set_keep_with_next(paragraph, heading)


def add_field(paragraph, instruction_text: str, fallback: str) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = f" {instruction_text} "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = fallback
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for node in (begin, instruction, separate, text, end):
        run._r.append(node)
    set_run_font(run, LATIN_FONT, 8.2, color=MUTED)


def configure_page(document: Document, *, header_label: str, footer_label: str) -> None:
    for section in document.sections:
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)
        section.header_distance = Cm(0.8)
        section.footer_distance = Cm(0.8)
        section.different_first_page_header_footer = False

        header = section.header.paragraphs[0]
        header.clear()
        header.alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_spacing(header, line=1.0)
        set_run_font(header.add_run(f"{header_label}  ·  "), LATIN_FONT, 8.2, color=MUTED)
        add_field(header, 'STYLEREF "Heading 2"', "Bilingual study edition")

        footer = section.footer.paragraphs[0]
        footer.clear()
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_spacing(footer, line=1.0)
        set_run_font(footer.add_run(f"{footer_label}  ·  "), CJK_FONT, 8.2, color=MUTED)
        add_field(footer, "PAGE", "1")

    document.settings.odd_and_even_pages_header_footer = False


def apply_styles(
    document: Document,
    *,
    document_title: str,
    header_label: str,
    footer_label: str,
) -> dict:
    configure_page(document, header_label=header_label, footer_label=footer_label)
    normal = document.styles["Normal"]
    normal.font.name = LATIN_FONT
    normal._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), LATIN_FONT)
    normal._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:hAnsi"), LATIN_FONT)
    normal._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), CJK_FONT)
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(INK)

    paragraphs = document.paragraphs
    callout_ranges, marker_indices = find_callout_ranges(paragraphs)
    callout_indices = {index for start, end, _ in callout_ranges for index in range(start, end + 1)}

    previous = None
    previous_heading_level: int | None = None
    for index, paragraph in enumerate(paragraphs):
        text = paragraph.text.strip()
        if index in callout_indices:
            previous = paragraph
            previous_heading_level = None
            continue
        if is_heading(paragraph):
            style_heading(paragraph)
            previous_heading_level = int(paragraph.style.name.rsplit(" ", 1)[-1])
            previous = paragraph
            continue
        if text and has_cjk(text) and previous is not None and (
            previous_heading_level is not None or (is_bold_label(previous) and is_bold_label(paragraph))
        ):
            style_heading_translation(paragraph, previous_heading_level)
            previous_heading_level = None
            previous = paragraph
            continue
        previous_heading_level = None
        if paragraph.style.name == "Source Code":
            style_code(paragraph)
        elif is_math_paragraph(paragraph):
            style_math(paragraph)
        elif not text:
            set_spacing(paragraph, after=2, line=1.0)
        elif has_cjk(text):
            style_chinese_body(paragraph)
        else:
            style_english_body(paragraph)
        previous = paragraph

    for start, end, role in callout_ranges:
        style_callout(paragraphs, start, end, role)

    for marker_index in reversed(marker_indices):
        marker = paragraphs[marker_index]._element
        marker.getparent().remove(marker)

    for table in document.tables:
        style_table(table)

    for index in range(1, len(paragraphs)):
        current = paragraphs[index]
        previous = paragraphs[index - 1]
        if (
            current.text.strip()
            and has_cjk(current.text)
            and index not in callout_indices
            and previous.text.strip()
            and not has_cjk(previous.text)
            and not is_heading(previous)
            and previous.style.name != "Source Code"
        ):
            set_keep_with_next(previous, True)

    document.core_properties.title = document_title
    document.core_properties.subject = "English-first Simplified-Chinese study edition"
    report = {
        "paragraph_count": len(paragraphs),
        "table_count": len(document.tables),
        "table_cell_count": sum(
            len(row.cells) for table in document.tables for row in table.rows
        ),
        "problem_callouts": sum(1 for _, _, role in callout_ranges if role == "problem"),
        "problem_numbered_paragraphs": sum(
            1
            for start, end, role in callout_ranges
            if role == "problem"
            for index in range(start, end + 1)
            if is_numbered_paragraph(paragraphs[index])
        ),
        "problem_numbering_origins_explicit": sum(
            1
            for start, end, role in callout_ranges
            if role == "problem"
            for index in range(start, end + 1)
            if is_numbered_paragraph(paragraphs[index])
            and paragraphs[index]._p.pPr.ind is not None
            and paragraphs[index]._p.pPr.ind.get(qn("w:firstLine")) == "0"
            and paragraphs[index]._p.pPr.ind.get(qn("w:hanging")) is None
        ),
        "problem_legacy_horizontal_rules": sum(
            1
            for start, end, role in callout_ranges
            if role == "problem"
            for index in range(start, end + 1)
            if has_legacy_horizontal_rule(paragraphs[index])
        ),
        "standalone_tip_callouts": sum(1 for _, _, role in callout_ranges if role == "tip"),
        "semantic_tip_labels": (
            sum(
                bool(CALLOUT_SOURCE_PATTERNS["tip"].search(paragraph.text.strip()))
                for paragraph in paragraphs
            )
            if "tip" in CALLOUT_SOURCE_PATTERNS
            else 0
        ),
        "example_callouts": sum(1 for _, _, role in callout_ranges if role == "example"),
    }
    if ACTIVE_SCHEMA_VERSION == 2:
        role_callouts: dict[str, int] = {role: 0 for role in ROLE_STYLES}
        occurrence_ids: dict[str, list[str]] = {role: [] for role in ROLE_STYLES}
        complete_role_callouts: dict[str, int] = {role: 0 for role in ROLE_STYLES}
        anchor_only_role_callouts: dict[str, int] = {role: 0 for role in ROLE_STYLES}
        occurrence_memberships: dict[str, str] = {}
        for item in callout_ranges:
            role_callouts[item[2]] += 1
            group_id = CALLOUT_RANGE_IDS[item]
            membership = CALLOUT_RANGE_MEMBERSHIPS[item]
            occurrence_ids[item[2]].append(group_id)
            occurrence_memberships[group_id] = membership
            target = (
                complete_role_callouts
                if membership == "complete"
                else anchor_only_role_callouts
            )
            target[item[2]] += 1
        report["role_callouts"] = role_callouts
        report["complete_role_callouts"] = complete_role_callouts
        report["anchor_only_role_callouts"] = anchor_only_role_callouts
        report["structural_occurrence_ids"] = occurrence_ids
        report["structural_occurrence_memberships"] = occurrence_memberships
    return report


def main() -> None:
    global LATIN_FONT, CJK_FONT, CODE_FONT
    parser = argparse.ArgumentParser()
    parser.add_argument("input", type=Path)
    parser.add_argument("output", type=Path)
    parser.add_argument("--title", default="English-Chinese Bilingual Study Edition")
    parser.add_argument("--header-label", default="Bilingual study edition")
    parser.add_argument("--footer-label", default="英中双语学习版")
    parser.add_argument("--latin-font", default=LATIN_FONT)
    parser.add_argument("--cjk-font", default=CJK_FONT)
    parser.add_argument("--code-font", default=CODE_FONT)
    parser.add_argument("--work-dir", type=Path)
    args = parser.parse_args()
    LATIN_FONT = args.latin_font
    CJK_FONT = args.cjk_font
    CODE_FONT = args.code_font
    if args.work_dir is None:
        work_dir = None
        input_path = args.input
        output_path = args.output
        document = Document(input_path)
    else:
        work_dir = lexical_absolute_path(args.work_dir)
        validate_artifact_directory(work_dir)
        stage_dir = work_dir / "output" / "docx-build"
        validate_artifact_tree(stage_dir, boundary=work_dir, allow_missing=False)
        input_path = _bounded_cli_path(
            stage_dir, args.input, label="DOCX style input"
        )
        output_path = _bounded_cli_path(
            stage_dir, args.output, label="DOCX style output"
        )
        input_key = tuple(part.casefold() for part in input_path.relative_to(stage_dir).parts)
        output_key = tuple(part.casefold() for part in output_path.relative_to(stage_dir).parts)
        if input_key == output_key:
            raise ArtifactSafetyError("DOCX style input and output roles must be distinct")
        validate_artifact_file(input_path, boundary=stage_dir)
        if os.path.lexists(output_path.parent):
            validate_artifact_file(
                output_path, boundary=stage_dir, allow_missing=True
            )
        document = Document(
            io.BytesIO(read_artifact_bytes(input_path, boundary=stage_dir))
        )
    report = apply_styles(
        document,
        document_title=args.title,
        header_label=args.header_label,
        footer_label=args.footer_label,
    )
    if work_dir is None:
        output_path.parent.mkdir(parents=True, exist_ok=True)
        document.save(output_path)
    else:
        prepare_artifact_directory(output_path.parent, boundary=stage_dir)
        atomic_publish_with_writer(
            output_path,
            document.save,
            boundary=stage_dir,
        )
    print(report)


if __name__ == "__main__":
    try:
        main()
    except ArtifactSafetyError as exc:
        raise SystemExit(str(exc)) from exc
