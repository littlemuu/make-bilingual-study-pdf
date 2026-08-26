#!/usr/bin/env python3
from __future__ import annotations

import json
import shutil
import subprocess
import sys
import tempfile
from pathlib import Path

import fitz
from PIL import Image
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from lxml import etree

from common import read_json, read_jsonl, sha256_file, sha256_text, write_json, write_jsonl
from build_outputs import latex_escape
from docx_style import (
    configure_profile,
    find_callout_ranges,
    has_legacy_horizontal_rule,
    style_callout,
)
from extract_pdf import invalid_pngs, repair_truncated_renders
from document_ir import migrate_work_dir, validate_ir_against_sources
from pipeline import translation_plan_status
from profile import canonical_profile_sha256, load_profile, semantic_match
from translation_utils import protect_source, restore_placeholders


SCRIPT_DIR = Path(__file__).resolve().parent


def working_command(name: str) -> str:
    """Return a probed executable, bypassing broken shell wrappers if needed."""
    seen: set[str] = set()
    for candidate in (shutil.which(name), shutil.which(f"{name}.exe")):
        if not candidate or candidate in seen:
            continue
        seen.add(candidate)
        try:
            probe = subprocess.run(
                [candidate, "-v"],
                check=False,
                stdout=subprocess.PIPE,
                stderr=subprocess.PIPE,
                text=True,
                timeout=10,
            )
        except (OSError, subprocess.TimeoutExpired):
            continue
        if probe.returncode == 0:
            return candidate
    raise AssertionError(f"no working {name} executable is available")


def run(script: str, work_dir: Path, expect_success: bool) -> dict:
    completed = subprocess.run(
        [sys.executable, str(SCRIPT_DIR / script), str(work_dir)],
        check=False,
        stdout=subprocess.PIPE,
        stderr=subprocess.STDOUT,
        text=True,
    )
    if (completed.returncode == 0) != expect_success:
        raise AssertionError(
            f"{script} returned {completed.returncode}, expected success={expect_success}\n"
            f"{completed.stdout}"
        )
    return read_json(work_dir / "translation" / "translation-audit.json")


def write_responses(work_dir: Path, mutate=None) -> None:
    translation_dir = work_dir / "translation"
    requests = read_jsonl(translation_dir / "requests" / "part-0001.jsonl")
    responses = []
    for request in requests:
        placeholders = " ".join(
            item["placeholder"] for item in request["protected_tokens"]
        )
        prefix = (
            "这是完整的翻译结果测试译文。"
            if request["id"] == "p001-b002"
            else "这是完整的测试译文。"
        )
        translation = f"{prefix}{placeholders}".strip()
        responses.append(
            {
                "id": request["id"],
                "source_sha256": request["source_sha256"],
                "translation": translation,
            }
        )
    if mutate:
        mutate(responses, requests)
    write_jsonl(translation_dir / "responses" / "part-0001.jsonl", responses)


def setup(work_dir: Path) -> None:
    work_dir.mkdir()
    profile = load_profile("assignment-en-zh")
    write_json(work_dir / "profile.json", profile)
    write_json(work_dir / "document-ir.json", {"fixture": True})
    blocks = [
        {
            "id": "p001-b001",
            "page": 1,
            "bbox": [10, 10, 300, 30],
            "source": "Use torch.Tensor with batch_size 32 in this complete example.",
            "kind": "prose",
            "translatable": True,
            "links": [],
            "protected_spans": [],
        },
        {
            "id": "p001-b002",
            "page": 1,
            "bbox": [10, 40, 300, 60],
            "source": (
                "This deliberately longer sentence checks that the translated result remains "
                "complete faithful accurate readable and useful for every careful student."
            ),
            "kind": "prose",
            "translatable": True,
            "links": [],
            "protected_spans": [],
        },
    ]
    for block in blocks:
        block["source_sha256"] = sha256_text(block["source"])
    write_jsonl(work_dir / "blocks.jsonl", blocks)
    write_json(
        work_dir / "manifest.json",
        {"source_sha256": "0" * 64, "problem_ids": [], "external_uris": []},
    )
    write_json(work_dir / "source-audit.json", {"status": "passed"})
    translation_dir = work_dir / "translation"
    translation_dir.mkdir()
    write_json(
        translation_dir / "glossary.json",
        {
            "schema_version": 1,
            "target_language": "zh-CN",
            "source_blocks_sha256": sha256_file(work_dir / "blocks.jsonl"),
            "terms": [
                {
                    "source": "translated result",
                    "targets": ["翻译结果"],
                    "case_sensitive": False,
                    "enforce": True,
                    "notes": "self-test",
                }
            ],
        },
    )
    subprocess.run(
        [
            sys.executable,
            str(SCRIPT_DIR / "prepare_translation.py"),
            str(work_dir),
            "--max-source-chars",
            "1000",
        ],
        check=True,
        stdout=subprocess.DEVNULL,
    )


def main() -> None:
    results = []
    profile = load_profile("assignment-en-zh")
    assert profile["translation"]["target_language"] == "zh-CN"
    assert semantic_match(profile, "Problem (profile_fixture): Test")["role"] == "problem"
    assert semantic_match(profile, "Example (profile_fixture): Test")["role"] == "example"
    assert semantic_match(profile, "Low-Resource Tip: Test")["role"] == "tip"
    results.append("default profile validates and classifies all assignment semantic roles")

    with tempfile.TemporaryDirectory(prefix="bilingual-ir-test-") as temp:
        work_dir = Path(temp)
        fixture_blocks = [
            {
                "id": "p001-b001",
                "page": 1,
                "bbox": [10, 10, 100, 30],
                "source": "Problem (ir_fixture): Test",
                "kind": "callout",
                "translatable": True,
                "protected_spans": [],
                "links": [],
            },
            {
                "id": "p001-b002",
                "page": 1,
                "bbox": [10, 40, 100, 60],
                "source": "Example (ir_example): Test",
                "kind": "callout",
                "translatable": True,
                "protected_spans": [],
                "links": [],
            },
            {
                "id": "p001-b003",
                "page": 1,
                "bbox": [10, 70, 100, 90],
                "source": "Low-Resource Tip: Test",
                "kind": "callout",
                "translatable": True,
                "protected_spans": [],
                "links": [],
            },
        ]
        for block in fixture_blocks:
            block["source_sha256"] = sha256_text(block["source"])
        write_jsonl(work_dir / "blocks.jsonl", fixture_blocks)
        write_json(
            work_dir / "manifest.json",
            {
                "schema_version": 3,
                "source_pdf": "/fixture/source.pdf",
                "source_sha256": "0" * 64,
                "page_count": 1,
                "external_uris": [],
                "visuals": [],
                "artifacts": {"blocks": "blocks.jsonl"},
            },
        )
        ir_path = migrate_work_dir(work_dir, "assignment-en-zh")
        ir = read_json(ir_path)
        assert ir["inventories"]["semantic_role_counts"] == {
            "problem": 1,
            "example": 1,
            "tip": 1,
        }
        assert all(
            group["membership"] == "anchor-only" for group in ir["semantic_groups"]
        )
        assert validate_ir_against_sources(work_dir) == []
        fixture_blocks[0]["source"] = "Problem (ir_fixture): Changed"
        fixture_blocks[0]["source_sha256"] = sha256_text(fixture_blocks[0]["source"])
        write_jsonl(work_dir / "blocks.jsonl", fixture_blocks)
        assert validate_ir_against_sources(work_dir)
        results.append("document IR is profile-bound, explicit about evidence, and detects drift")
    wrapped_url_block = {
        "source": "Read https://example.\ncom/study_fixture for details.",
        "protected_spans": [
            {"start": 5, "end": 21, "role": "code"},
            {"start": 22, "end": 39, "role": "code"},
        ],
    }
    protected_text, protected_tokens = protect_source(wrapped_url_block)
    assert len(protected_tokens) == 1 and protected_tokens[0]["role"] == "url"
    assert protected_tokens[0]["value"] == "https://example.com/study_fixture"
    assert restore_placeholders(protected_text, protected_tokens) == (
        "Read https://example.com/study_fixture for details."
    )
    results.append("wrapped URL becomes one canonical placeholder")

    expected_math_replacements = {
        "‣": r"\BilingualMath{\blacktriangleright}",
        "▷": r"\BilingualMath{\triangleright}",
        "≤": r"\BilingualMath{\leq}",
        "≥": r"\BilingualMath{\geq}",
        "∈": r"\BilingualMath{\in}",
        "ℝ": r"\BilingualMath{\mathbb{R}}",
        "⊤": r"\BilingualMath{\top}",
        "𝟙": r"\BilingualMath{\mathbb{1}}",
        "′": r"\BilingualMath{^{\prime}}",
        "⋅": r"\BilingualMath{\cdot}",
        "∇": r"\BilingualMath{\nabla}",
        "⊙": r"\BilingualMath{\odot}",
        "∑": r"\BilingualMath{\sum}",
        "𝛼": r"\BilingualMath{𝛼}",
        "𝛽": r"\BilingualMath{𝛽}",
        "𝜀": r"\BilingualMath{𝜀}",
        "𝜇": r"\BilingualMath{𝜇}",
        "𝜎": r"\BilingualMath{𝜎}",
        "𝜃": r"\BilingualMath{𝜃}",
        "𝜆": r"\BilingualMath{𝜆}",
        "𝜏": r"\BilingualMath{𝜏}",
        "ℎ": r"\BilingualMath{ℎ}",
        "ℓ": r"\BilingualMath{\ell}",
    }
    for source, expected in expected_math_replacements.items():
        assert latex_escape(source) == expected
    assert latex_escape("𝒩︀") == r"\BilingualMath{\mathcal{N}}"
    assert latex_escape("︀") == "︀"
    results.append("Unicode math symbols use distinct lossless LaTeX replacements")

    italic_pairs = (
        [(chr(codepoint), letter) for codepoint, letter in zip(
            range(0x1D434, 0x1D44E), "ABCDEFGHIJKLMNOPQRSTUVWXYZ"
        )]
        + [(chr(codepoint), letter) for codepoint, letter in zip(
            range(0x1D44E, 0x1D455), "abcdefg"
        )]
        + [("ℎ", "h")]
        + [(chr(codepoint), letter) for codepoint, letter in zip(
            range(0x1D456, 0x1D468), "ijklmnopqrstuvwxyz"
        )]
    )
    assert len(italic_pairs) == 52
    for source, _letter in italic_pairs:
        assert latex_escape(source) == rf"\BilingualMath{{{source}}}"
    results.append("all mathematical italic Latin letters retain their Unicode identity")

    template = (SCRIPT_DIR.parent / "assets" / "bilingual-template.tex").read_text(
        encoding="utf-8"
    )
    assert r"\usepackage{amssymb}" in template
    assert r"\usepackage{unicode-math}" in template
    assert r"\setmathfont{Latin Modern Math}" in template
    assert r"\DeclareRobustCommand{\BilingualMath}[1]{\ensuremath{#1}}" in template
    assert r"\RenewDocumentCommand{\triangleright}" not in template
    results.append("template provides Unicode math without redefining standard commands")

    with tempfile.TemporaryDirectory(prefix="bilingual-png-integrity-test-") as temp:
        temp_dir = Path(temp)
        valid_png = temp_dir / "valid.png"
        truncated_png = temp_dir / "truncated.png"
        Image.new("RGB", (64, 64), "white").save(valid_png)
        payload = valid_png.read_bytes()
        truncated_png.write_bytes(payload[: len(payload) // 2])
        assert invalid_pngs([valid_png, truncated_png]) == [truncated_png]
        results.append("fully decoded PNG audit detects truncated renders")

    with tempfile.TemporaryDirectory(prefix="bilingual-png-repair-test-") as temp:
        temp_dir = Path(temp)
        pdf_path = temp_dir / "fixture.pdf"
        document = fitz.open()
        page = document.new_page()
        page.insert_text((72, 72), "render repair fixture")
        document.save(pdf_path)
        document.close()
        broken_render = temp_dir / "page-1.png"
        broken_render.write_bytes(b"not a PNG")
        pdftoppm = working_command("pdftoppm")
        repair_truncated_renders(pdftoppm, pdf_path, [broken_render], 72)
        assert invalid_pngs([broken_render]) == []
        results.append("single-page retry repairs a truncated PNG with a blocking copy")

    with tempfile.TemporaryDirectory(prefix="bilingual-problem-border-test-") as temp:
        temp_dir = Path(temp)
        document = Document()
        paragraphs = [
            document.add_paragraph("Problem (border_fixture): Border fixture"),
            document.add_paragraph("First numbered requirement."),
            document.add_paragraph("Deliverable: A short response."),
            document.add_paragraph(),
            document.add_paragraph("问题（border_fixture）：边框测试"),
            document.add_paragraph("第一条编号要求。"),
            document.add_paragraph("交付物：简短回答。"),
        ]
        for paragraph in (paragraphs[1], paragraphs[5]):
            p_pr = paragraph._p.get_or_add_pPr()
            num_pr = OxmlElement("w:numPr")
            level = OxmlElement("w:ilvl")
            level.set(qn("w:val"), "0")
            num_id = OxmlElement("w:numId")
            num_id.set(qn("w:val"), "1")
            num_pr.extend((level, num_id))
            p_pr.append(num_pr)
        pict = OxmlElement("w:pict")
        rect = etree.Element(
            "{urn:schemas-microsoft-com:vml}rect",
            nsmap={
                "v": "urn:schemas-microsoft-com:vml",
                "o": "urn:schemas-microsoft-com:office:office",
            },
        )
        rect.set("{urn:schemas-microsoft-com:office:office}hr", "t")
        pict.append(rect)
        paragraphs[3].add_run()._r.append(pict)

        style_callout(paragraphs, 0, len(paragraphs) - 1, "problem")
        numbered = (paragraphs[1], paragraphs[5])
        for paragraph in numbered:
            ind = paragraph._p.pPr.ind
            assert ind is not None
            assert ind.get(qn("w:firstLine")) == "0"
            assert ind.get(qn("w:hanging")) is None
        assert not any(has_legacy_horizontal_rule(paragraph) for paragraph in paragraphs)
        indents = {
            (
                paragraph._p.pPr.ind.get(qn("w:left")),
                paragraph._p.pPr.ind.get(qn("w:right")),
            )
            for paragraph in paragraphs
        }
        assert len(indents) == 1
        fixture_docx = temp_dir / "problem-border-fixture.docx"
        document.save(fixture_docx)
        audit = subprocess.run(
            [
                sys.executable,
                str(SCRIPT_DIR / "audit_docx.py"),
                str(fixture_docx),
                "--expected-problems",
                "1",
            ],
            check=False,
            capture_output=True,
            text=True,
        )
        assert audit.returncode == 0, audit.stdout + audit.stderr
        results.append("numbered Problem borders share one origin and one bounded divider")

    configure_profile(profile)
    document = Document()
    document.styles.add_style("Block Text", WD_STYLE_TYPE.PARAGRAPH)
    callout_paragraphs = [
        document.add_paragraph("Low-Resource Tip: Source label", style="Block Text"),
        document.add_paragraph("低资源提示：目标语言标签", style="Block Text"),
        document.add_paragraph("Example (range_fixture): Source label", style="Block Text"),
        document.add_paragraph("示例（range_fixture）：目标语言标签", style="Block Text"),
    ]
    ranges, _markers = find_callout_ranges(callout_paragraphs)
    assert [(start, end, role) for start, end, role in ranges] == [
        (0, 1, "tip"),
        (2, 3, "example"),
    ]
    results.append("target-language labels stay inside rather than reopening semantic callouts")

    with tempfile.TemporaryDirectory(prefix="bilingual-docx-finalize-test-") as temp:
        work_dir = Path(temp)
        output_dir = work_dir / "output"
        translation_dir = work_dir / "translation"
        output_dir.mkdir()
        translation_dir.mkdir()
        write_json(work_dir / "manifest.json", {"source_sha256": "0" * 64})
        for path, payload in (
            (output_dir / "fixture.md", "source\n\n译文\n"),
            (output_dir / "fixture.tex", "fixture\n"),
            (output_dir / "fixture.docx", "docx fixture\n"),
            (output_dir / "fixture.pdf", "pdf fixture\n"),
        ):
            path.write_text(payload, encoding="utf-8")
        write_json(
            output_dir / "build-manifest.json",
            {"markdown": "fixture.md", "latex": "fixture.tex"},
        )
        write_json(work_dir / "source-audit.json", {"status": "passed"})
        write_json(translation_dir / "translation-audit.json", {"status": "passed"})
        write_json(output_dir / "output-audit.json", {"status": "passed"})
        pdf_hash = sha256_file(output_dir / "fixture.pdf")
        contact_dir = output_dir / "contact"
        contact_dir.mkdir()
        contact_path = contact_dir / "contact-001.png"
        Image.new("RGB", (64, 64), "white").save(contact_path)
        contact_record = {
            "path": "contact/contact-001.png",
            "sha256": sha256_file(contact_path),
            "first_page": 1,
            "last_page": 1,
        }
        write_json(
            output_dir / "compile-audit.json",
            {
                "status": "passed",
                "automated_status": "passed",
                "docx": "fixture.docx",
                "pdf": "fixture.pdf",
                "pdf_sha256": pdf_hash,
                "page_count": 1,
                "contact_sheets": [contact_record],
                "warnings": [],
            },
        )
        write_json(
            output_dir / "visual-review.json",
            {
                "status": "passed",
                "pdf_sha256": pdf_hash,
                "reviewed_pages": [1],
                "contact_sheets_inspected": [contact_record["path"]],
                "contact_sheets_sha256": {
                    contact_record["path"]: contact_record["sha256"]
                },
                "notes": "Inspected the synthetic one-page contact sheet.",
            },
        )
        finalized = subprocess.run(
            [sys.executable, str(SCRIPT_DIR / "finalize_qa.py"), str(work_dir)],
            check=False,
            capture_output=True,
            text=True,
        )
        assert finalized.returncode == 0, finalized.stdout + finalized.stderr
        qa = read_json(output_dir / "qa-report.json")
        assert qa["status"] == "passed" and "docx" in qa["deliverables"]
        results.append("DOCX-profile compile and visual gates finalize all editable deliverables")

    with tempfile.TemporaryDirectory(prefix="bilingual-skill-self-test-") as temp:
        work_dir = Path(temp) / "work"
        setup(work_dir)

        plan_path = work_dir / "translation" / "plan.json"
        plan = read_json(plan_path)
        assert plan["profile_sha256"] == canonical_profile_sha256(profile)
        assert plan["profile_file_sha256"] == sha256_file(work_dir / "profile.json")
        assert translation_plan_status(work_dir, plan_path) == "passed"
        request_path = work_dir / "translation" / plan["batches"][0]["request_file"]
        request_path.write_text(
            request_path.read_text(encoding="utf-8") + "\n", encoding="utf-8"
        )
        assert translation_plan_status(work_dir, plan_path) == "stale"
        subprocess.run(
            [
                sys.executable,
                str(SCRIPT_DIR / "prepare_translation.py"),
                str(work_dir),
                "--max-source-chars",
                "1000",
                "--force",
            ],
            check=True,
            stdout=subprocess.DEVNULL,
        )
        results.append("canonical Profile hashes and frozen request batches have distinct verified bindings")

        write_responses(work_dir)
        report = run("audit_translation.py", work_dir, True)
        assert report["status"] == "passed"
        results.append("valid response passes")

        def missing(responses, _requests):
            responses.pop()

        write_responses(work_dir, missing)
        report = run("audit_translation.py", work_dir, False)
        assert report["status"] == "incomplete" and report["missing_ids"]
        results.append("missing response is incomplete")

        def duplicate(responses, _requests):
            responses.append(dict(responses[0]))

        write_responses(work_dir, duplicate)
        report = run("audit_translation.py", work_dir, False)
        assert report["duplicate_ids"]
        results.append("duplicate response fails")

        def stale_hash(responses, _requests):
            responses[0]["source_sha256"] = "f" * 64

        write_responses(work_dir, stale_hash)
        report = run("audit_translation.py", work_dir, False)
        assert report["invalid_source_hash_ids"]
        results.append("stale source hash fails")

        def broken_placeholder(responses, requests):
            placeholder = requests[0]["protected_tokens"][0]["placeholder"]
            responses[0]["translation"] = responses[0]["translation"].replace(
                placeholder, "⟦BROKEN⟧"
            )

        write_responses(work_dir, broken_placeholder)
        report = run("audit_translation.py", work_dir, False)
        assert report["placeholder_failures"]
        results.append("changed placeholder fails")

        def missing_glossary_target(responses, _requests):
            responses[1]["translation"] = "这是完整的中文测试译文。"

        write_responses(work_dir, missing_glossary_target)
        report = run("audit_translation.py", work_dir, False)
        assert report["glossary_failures"]
        results.append("enforced glossary omission fails")

        def source_copy(responses, requests):
            responses[1]["translation"] = "译文：" + requests[1]["source_for_translation"]

        write_responses(work_dir, source_copy)
        report = run("audit_translation.py", work_dir, False)
        assert report["source_copy_ids"]
        results.append("substantially copied English fails")

    print(json.dumps({"status": "passed", "tests": results}, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
