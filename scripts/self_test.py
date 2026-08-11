#!/usr/bin/env python3
from __future__ import annotations

import json
import shutil
import subprocess
import sys
import tempfile
from pathlib import Path

import fitz
from PIL import Image
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from lxml import etree

from common import read_json, read_jsonl, sha256_file, sha256_text, write_json, write_jsonl
from build_outputs import latex_escape
from docx_style import has_legacy_horizontal_rule, style_callout
from extract_pdf import invalid_pngs, repair_truncated_renders
from translation_utils import protect_source, restore_placeholders


SCRIPT_DIR = Path(__file__).resolve().parent


def run(script: str, work_dir: Path, expect_success: bool) -> dict:
    completed = subprocess.run(
        [sys.executable, str(SCRIPT_DIR / script), str(work_dir)],
        check=False,
        stdout=subprocess.PIPE,
        stderr=subprocess.STDOUT,
        text=True,
    )
    if (completed.returncode == 0) != expect_success:
        raise AssertionError(
            f"{script} returned {completed.returncode}, expected success={expect_success}\n"
            f"{completed.stdout}"
        )
    return read_json(work_dir / "translation" / "translation-audit.json")


def write_responses(work_dir: Path, mutate=None) -> None:
    translation_dir = work_dir / "translation"
    requests = read_jsonl(translation_dir / "requests" / "part-0001.jsonl")
    responses = []
    for request in requests:
        placeholders = " ".join(
            item["placeholder"] for item in request["protected_tokens"]
        )
        prefix = (
            "这是完整的翻译结果测试译文。"
            if request["id"] == "p001-b002"
            else "这是完整的测试译文。"
        )
        translation = f"{prefix}{placeholders}".strip()
        responses.append(
            {
                "id": request["id"],
                "source_sha256": request["source_sha256"],
                "translation": translation,
            }
        )
    if mutate:
        mutate(responses, requests)
    write_jsonl(translation_dir / "responses" / "part-0001.jsonl", responses)


def setup(work_dir: Path) -> None:
    work_dir.mkdir()
    blocks = [
        {
            "id": "p001-b001",
            "page": 1,
            "bbox": [10, 10, 300, 30],
            "source": "Use torch.Tensor with batch_size 32 in this complete example.",
            "kind": "prose",
            "translatable": True,
            "links": [],
            "protected_spans": [],
        },
        {
            "id": "p001-b002",
            "page": 1,
            "bbox": [10, 40, 300, 60],
            "source": (
                "This deliberately longer sentence checks that the translated result remains "
                "complete faithful accurate readable and useful for every careful student."
            ),
            "kind": "prose",
            "translatable": True,
            "links": [],
            "protected_spans": [],
        },
    ]
    for block in blocks:
        block["source_sha256"] = sha256_text(block["source"])
    write_jsonl(work_dir / "blocks.jsonl", blocks)
    write_json(
        work_dir / "manifest.json",
        {"source_sha256": "0" * 64, "problem_ids": [], "external_uris": []},
    )
    write_json(work_dir / "source-audit.json", {"status": "passed"})
    translation_dir = work_dir / "translation"
    translation_dir.mkdir()
    write_json(
        translation_dir / "glossary.json",
        {
            "schema_version": 1,
            "target_language": "zh-CN",
            "source_blocks_sha256": sha256_file(work_dir / "blocks.jsonl"),
            "terms": [
                {
                    "source": "translated result",
                    "targets": ["翻译结果"],
                    "case_sensitive": False,
                    "enforce": True,
                    "notes": "self-test",
                }
            ],
        },
    )
    subprocess.run(
        [
            sys.executable,
            str(SCRIPT_DIR / "prepare_translation.py"),
            str(work_dir),
            "--max-source-chars",
            "1000",
        ],
        check=True,
        stdout=subprocess.DEVNULL,
    )


def main() -> None:
    results = []
    wrapped_url_block = {
        "source": "Read https://example.\ncom/study_fixture for details.",
        "protected_spans": [
            {"start": 5, "end": 21, "role": "code"},
            {"start": 22, "end": 39, "role": "code"},
        ],
    }
    protected_text, protected_tokens = protect_source(wrapped_url_block)
    assert len(protected_tokens) == 1 and protected_tokens[0]["role"] == "url"
    assert protected_tokens[0]["value"] == "https://example.com/study_fixture"
    assert restore_placeholders(protected_text, protected_tokens) == (
        "Read https://example.com/study_fixture for details."
    )
    results.append("wrapped URL becomes one canonical placeholder")

    expected_math_replacements = {
        "‣": r"\BilingualMath{\blacktriangleright}",
        "▷": r"\BilingualMath{\triangleright}",
        "≤": r"\BilingualMath{\leq}",
        "≥": r"\BilingualMath{\geq}",
        "∈": r"\BilingualMath{\in}",
        "ℝ": r"\BilingualMath{\mathbb{R}}",
        "⊤": r"\BilingualMath{\top}",
        "𝟙": r"\BilingualMath{\mathbb{1}}",
        "′": r"\BilingualMath{^{\prime}}",
        "⋅": r"\BilingualMath{\cdot}",
        "∇": r"\BilingualMath{\nabla}",
        "⊙": r"\BilingualMath{\odot}",
        "∑": r"\BilingualMath{\sum}",
        "𝛼": r"\BilingualMath{𝛼}",
        "𝛽": r"\BilingualMath{𝛽}",
        "𝜀": r"\BilingualMath{𝜀}",
        "𝜇": r"\BilingualMath{𝜇}",
        "𝜎": r"\BilingualMath{𝜎}",
        "𝜃": r"\BilingualMath{𝜃}",
        "𝜆": r"\BilingualMath{𝜆}",
        "𝜏": r"\BilingualMath{𝜏}",
        "ℎ": r"\BilingualMath{ℎ}",
        "ℓ": r"\BilingualMath{\ell}",
    }
    for source, expected in expected_math_replacements.items():
        assert latex_escape(source) == expected
    assert latex_escape("𝒩︀") == r"\BilingualMath{\mathcal{N}}"
    assert latex_escape("︀") == "︀"
    results.append("Unicode math symbols use distinct lossless LaTeX replacements")

    italic_pairs = (
        [(chr(codepoint), letter) for codepoint, letter in zip(
            range(0x1D434, 0x1D44E), "ABCDEFGHIJKLMNOPQRSTUVWXYZ"
        )]
        + [(chr(codepoint), letter) for codepoint, letter in zip(
            range(0x1D44E, 0x1D455), "abcdefg"
        )]
        + [("ℎ", "h")]
        + [(chr(codepoint), letter) for codepoint, letter in zip(
            range(0x1D456, 0x1D468), "ijklmnopqrstuvwxyz"
        )]
    )
    assert len(italic_pairs) == 52
    for source, _letter in italic_pairs:
        assert latex_escape(source) == rf"\BilingualMath{{{source}}}"
    results.append("all mathematical italic Latin letters retain their Unicode identity")

    template = (SCRIPT_DIR.parent / "assets" / "bilingual-template.tex").read_text(
        encoding="utf-8"
    )
    assert r"\usepackage{amssymb}" in template
    assert r"\usepackage{unicode-math}" in template
    assert r"\setmathfont{Latin Modern Math}" in template
    assert r"\DeclareRobustCommand{\BilingualMath}[1]{\ensuremath{#1}}" in template
    assert r"\RenewDocumentCommand{\triangleright}" not in template
    results.append("template provides Unicode math without redefining standard commands")

    with tempfile.TemporaryDirectory(prefix="bilingual-png-integrity-test-") as temp:
        temp_dir = Path(temp)
        valid_png = temp_dir / "valid.png"
        truncated_png = temp_dir / "truncated.png"
        Image.new("RGB", (64, 64), "white").save(valid_png)
        payload = valid_png.read_bytes()
        truncated_png.write_bytes(payload[: len(payload) // 2])
        assert invalid_pngs([valid_png, truncated_png]) == [truncated_png]
        results.append("fully decoded PNG audit detects truncated renders")

    with tempfile.TemporaryDirectory(prefix="bilingual-png-repair-test-") as temp:
        temp_dir = Path(temp)
        pdf_path = temp_dir / "fixture.pdf"
        document = fitz.open()
        page = document.new_page()
        page.insert_text((72, 72), "render repair fixture")
        document.save(pdf_path)
        document.close()
        broken_render = temp_dir / "page-1.png"
        broken_render.write_bytes(b"not a PNG")
        pdftoppm = shutil.which("pdftoppm")
        assert pdftoppm is not None
        repair_truncated_renders(pdftoppm, pdf_path, [broken_render], 72)
        assert invalid_pngs([broken_render]) == []
        results.append("single-page retry repairs a truncated PNG with a blocking copy")

    with tempfile.TemporaryDirectory(prefix="bilingual-problem-border-test-") as temp:
        temp_dir = Path(temp)
        document = Document()
        paragraphs = [
            document.add_paragraph("Problem (border_fixture): Border fixture"),
            document.add_paragraph("First numbered requirement."),
            document.add_paragraph("Deliverable: A short response."),
            document.add_paragraph(),
            document.add_paragraph("问题（border_fixture）：边框测试"),
            document.add_paragraph("第一条编号要求。"),
            document.add_paragraph("交付物：简短回答。"),
        ]
        for paragraph in (paragraphs[1], paragraphs[5]):
            p_pr = paragraph._p.get_or_add_pPr()
            num_pr = OxmlElement("w:numPr")
            level = OxmlElement("w:ilvl")
            level.set(qn("w:val"), "0")
            num_id = OxmlElement("w:numId")
            num_id.set(qn("w:val"), "1")
            num_pr.extend((level, num_id))
            p_pr.append(num_pr)
        pict = OxmlElement("w:pict")
        rect = etree.Element(
            "{urn:schemas-microsoft-com:vml}rect",
            nsmap={
                "v": "urn:schemas-microsoft-com:vml",
                "o": "urn:schemas-microsoft-com:office:office",
            },
        )
        rect.set("{urn:schemas-microsoft-com:office:office}hr", "t")
        pict.append(rect)
        paragraphs[3].add_run()._r.append(pict)

        style_callout(paragraphs, 0, len(paragraphs) - 1, "problem")
        numbered = (paragraphs[1], paragraphs[5])
        for paragraph in numbered:
            ind = paragraph._p.pPr.ind
            assert ind is not None
            assert ind.get(qn("w:firstLine")) == "0"
            assert ind.get(qn("w:hanging")) is None
        assert not any(has_legacy_horizontal_rule(paragraph) for paragraph in paragraphs)
        indents = {
            (
                paragraph._p.pPr.ind.get(qn("w:left")),
                paragraph._p.pPr.ind.get(qn("w:right")),
            )
            for paragraph in paragraphs
        }
        assert len(indents) == 1
        fixture_docx = temp_dir / "problem-border-fixture.docx"
        document.save(fixture_docx)
        audit = subprocess.run(
            [
                sys.executable,
                str(SCRIPT_DIR / "audit_docx.py"),
                str(fixture_docx),
                "--expected-problems",
                "1",
            ],
            check=False,
            capture_output=True,
            text=True,
        )
        assert audit.returncode == 0, audit.stdout + audit.stderr
        results.append("numbered Problem borders share one origin and one bounded divider")

    with tempfile.TemporaryDirectory(prefix="bilingual-skill-self-test-") as temp:
        work_dir = Path(temp) / "work"
        setup(work_dir)

        write_responses(work_dir)
        report = run("audit_translation.py", work_dir, True)
        assert report["status"] == "passed"
        results.append("valid response passes")

        def missing(responses, _requests):
            responses.pop()

        write_responses(work_dir, missing)
        report = run("audit_translation.py", work_dir, False)
        assert report["status"] == "incomplete" and report["missing_ids"]
        results.append("missing response is incomplete")

        def duplicate(responses, _requests):
            responses.append(dict(responses[0]))

        write_responses(work_dir, duplicate)
        report = run("audit_translation.py", work_dir, False)
        assert report["duplicate_ids"]
        results.append("duplicate response fails")

        def stale_hash(responses, _requests):
            responses[0]["source_sha256"] = "f" * 64

        write_responses(work_dir, stale_hash)
        report = run("audit_translation.py", work_dir, False)
        assert report["invalid_source_hash_ids"]
        results.append("stale source hash fails")

        def broken_placeholder(responses, requests):
            placeholder = requests[0]["protected_tokens"][0]["placeholder"]
            responses[0]["translation"] = responses[0]["translation"].replace(
                placeholder, "⟦BROKEN⟧"
            )

        write_responses(work_dir, broken_placeholder)
        report = run("audit_translation.py", work_dir, False)
        assert report["placeholder_failures"]
        results.append("changed placeholder fails")

        def missing_glossary_target(responses, _requests):
            responses[1]["translation"] = "这是完整的中文测试译文。"

        write_responses(work_dir, missing_glossary_target)
        report = run("audit_translation.py", work_dir, False)
        assert report["glossary_failures"]
        results.append("enforced glossary omission fails")

        def source_copy(responses, requests):
            responses[1]["translation"] = "译文：" + requests[1]["source_for_translation"]

        write_responses(work_dir, source_copy)
        report = run("audit_translation.py", work_dir, False)
        assert report["source_copy_ids"]
        results.append("substantially copied English fails")

    print(json.dumps({"status": "passed", "tests": results}, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
