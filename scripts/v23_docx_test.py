#!/usr/bin/env python3
"""Focused V2.3 DOCX semantic-container and frozen-audit regression tests."""
from __future__ import annotations

import json
import shutil
import subprocess
import sys
import tempfile
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE

import docx_style
import compile_docx_pdf
from build_docx import materialize_html_tables
from common import sha256_file, write_json
from docx_ast import GENERIC_BEGIN, GENERIC_END, PROBLEM_BEGIN, PROBLEM_END, transform
from profile import canonical_profile_sha256, load_profile, profile_contract


SCRIPT_DIR = Path(__file__).resolve().parent


def paragraph(text: str) -> dict:
    return {"t": "Para", "c": [{"t": "Str", "c": text}]}


def marker(node_id: str) -> dict:
    return {
        "t": "RawBlock",
        "c": ["html", f"<!-- bilingual:segment id={node_id} source_sha256=test -->"],
    }


def target(text: str) -> dict:
    return {"t": "BlockQuote", "c": [paragraph(text)]}


def test_legacy_transform() -> None:
    profile = load_profile("assignment-en-zh")
    document = {
        "pandoc-api-version": [1, 23],
        "meta": {},
        "blocks": [
            {
                "t": "BlockQuote",
                "c": [paragraph("Problem (legacy) source"), paragraph("问题（legacy）目标")],
            }
        ],
    }
    result = transform(document, profile)
    grouped = result["blocks"][0]["c"]
    assert grouped[0] == paragraph(PROBLEM_BEGIN)
    assert grouped[-1] == paragraph(PROBLEM_END)
    assert sum(block.get("t") == "HorizontalRule" for block in grouped) == 1
    assert result["meta"]["v2-problem-group-count"]["c"] == "1"


def test_generic_ast() -> None:
    profile = load_profile("lecture-notes-en-zh")
    document = {
        "pandoc-api-version": [1, 23],
        "meta": {},
        "blocks": [
            marker("a"),
            paragraph("Theorem 1. Source anchor"),
            target("定理 1。目标锚点"),
            marker("b"),
            paragraph("Second source member"),
            target("第二个目标成员"),
            marker("c"),
            paragraph("Lemma 2. Anchor only"),
            target("引理 2。仅锚点"),
            marker("d"),
            paragraph("Ordinary neighboring paragraph"),
            target("普通相邻段落"),
        ],
    }
    groups = [
        {
            "id": "theorem:a",
            "role": "theorem",
            "member_node_ids": ["a", "b"],
            "membership": "complete",
        },
        {
            "id": "lemma:c",
            "role": "lemma",
            "anchor_node_id": "c",
            "member_node_ids": ["c"],
            "membership": "anchor-only",
        },
        {
            "id": "paragraph:d",
            "role": "paragraph",
            "anchor_node_id": "d",
            "member_node_ids": ["d"],
            "membership": "anchor-only",
        },
    ]
    result = transform(document, profile, semantic_groups=groups)
    callouts = [block for block in result["blocks"] if block.get("t") == "BlockQuote"]
    assert len(callouts) == 3
    structural = callouts[0]["c"]
    assert GENERIC_BEGIN in structural[0]["c"][0]["c"]
    assert GENERIC_END in structural[-1]["c"][0]["c"]
    assert sum(block.get("t") == "HorizontalRule" for block in structural) == 1
    structural_text = " ".join(
        str(inline.get("c", ""))
        for block in structural
        for inline in block.get("c", [])
        if isinstance(inline, dict)
    )
    assert structural_text.index("Theorem 1. Source anchor") < structural_text.index("Second source member")
    assert structural_text.index("Second source member") < structural_text.index("定理 1。目标锚点")
    scoped_text = " ".join(
        str(inline.get("c", ""))
        for block in callouts[1]["c"]
        for inline in block.get("c", [])
        if isinstance(inline, dict)
    )
    assert "Lemma 2. Anchor only" in scoped_text
    assert "Ordinary neighboring paragraph" not in scoped_text
    assert any(block.get("t") == "RawBlock" for block in result["blocks"])
    counts = json.loads(result["meta"]["v23-structural-group-counts"]["c"])
    complete_counts = json.loads(
        result["meta"]["v23-complete-structural-group-counts"]["c"]
    )
    scoped_counts = json.loads(
        result["meta"]["v23-anchor-only-callout-counts"]["c"]
    )
    assert counts["theorem"] == 1 and counts["lemma"] == 1
    assert complete_counts["theorem"] == 1 and complete_counts["lemma"] == 0
    assert scoped_counts["theorem"] == 0 and scoped_counts["lemma"] == 1
    assert counts["paragraph"] == 0

    grouped_alias_document = {
        "pandoc-api-version": [1, 23],
        "meta": {},
        "blocks": [
            marker("b"),
            marker("a"),
            paragraph("Theorem 1. Source anchor Second source member"),
            target("定理一。目标锚点 第二个目标成员"),
        ],
    }
    grouped_alias = transform(
        grouped_alias_document, profile, semantic_groups=groups[:1]
    )
    assert len(grouped_alias["blocks"]) == 1
    assert grouped_alias["blocks"][0]["t"] == "BlockQuote"


def add_generic_range(
    document: Document,
    role: str,
    style: str,
    group_id: str,
    source: str,
    target_text: str,
    *,
    membership: str = "complete",
) -> None:
    document.add_paragraph(
        f"{GENERIC_BEGIN}|{role}|{style}|{membership}|{group_id}",
        style="Block Text",
    )
    document.add_paragraph(source, style="Block Text")
    document.add_paragraph("", style="Block Text")
    document.add_paragraph(target_text, style="Block Text")
    document.add_paragraph(
        f"{GENERIC_END}|{role}|{style}|{membership}|{group_id}",
        style="Block Text",
    )


def test_shared_style_roles() -> None:
    profile = load_profile("lecture-notes-en-zh")
    document = Document()
    document.styles.add_style("Block Text", WD_STYLE_TYPE.PARAGRAPH)
    add_generic_range(document, "theorem", "theorem", "theorem:a", "Theorem 1. A", "定理一")
    add_generic_range(
        document,
        "lemma",
        "theorem",
        "lemma:b",
        "Lemma 2. B",
        "引理二",
        membership="anchor-only",
    )
    document.add_paragraph("Section 1 is not a callout")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Metric"
    table.cell(0, 1).text = "Value"
    table.cell(1, 0).text = "Coverage"
    table.cell(1, 1).text = "100%"
    docx_style.configure_profile(profile)
    report = docx_style.apply_styles(
        document,
        document_title="V2.3 test",
        header_label="test",
        footer_label="测试",
    )
    assert report["role_callouts"]["theorem"] == 1
    assert report["role_callouts"]["lemma"] == 1
    assert report["complete_role_callouts"]["theorem"] == 1
    assert report["anchor_only_role_callouts"]["lemma"] == 1
    assert report["structural_occurrence_memberships"] == {
        "theorem:a": "complete",
        "lemma:b": "anchor-only",
    }
    assert report["structural_occurrence_ids"]["theorem"] == ["theorem:a"]
    assert report["structural_occurrence_ids"]["lemma"] == ["lemma:b"]
    assert report["table_count"] == 1
    assert report["table_cell_count"] == 4
    assert table.cell(0, 0)._tc.xpath("./w:tcPr/w:shd/@w:fill") == [
        docx_style.ACCENT_LIGHT
    ]
    assert table._tbl.xpath("./w:tblPr/w:tblBorders/w:insideV/@w:val") == [
        "single"
    ]
    assert all(GENERIC_BEGIN not in item.text and GENERIC_END not in item.text for item in document.paragraphs)


def test_html_table_materialization() -> None:
    if not shutil.which("pandoc"):
        return
    table_html = (
        "<table><thead><tr><th>Metric</th><th>Value</th></tr></thead>"
        "<tbody><tr><td>Coverage</td><td>100%</td></tr></tbody></table>"
    )
    parsed = subprocess.run(
        ["pandoc", "--from", "markdown", "--to", "json"],
        input=f"```{{=html}}\n{table_html}\n```\n",
        text=True,
        encoding="utf-8",
        capture_output=True,
        check=True,
    )
    ast = json.loads(parsed.stdout)
    assert ast["blocks"] == [{"t": "RawBlock", "c": ["html", table_html]}]
    materialized = materialize_html_tables(ast)
    assert [block["t"] for block in materialized["blocks"]] == ["Table"]

    reference_node = {
        "type": "list",
        "source": {
            "text": "- [1] Alpha (2025).\n- [2] Beta (2026).",
        },
    }
    assert "• [1] Alpha (2025). • [2] Beta (2026)." in (
        compile_docx_pdf.searchable_sources(reference_node)
    )


def make_inventory(profile: dict) -> dict:
    result = {}
    node_ids = {
        "theorem": "theorem-node",
        "definition": "definition-node",
        "equation": "equation-node",
    }
    for role, policy in profile_contract(profile)["role_inventory"].items():
        occurrence = 1 if role in node_ids else 0
        membership = "complete" if role == "theorem" else "anchor-only"
        result[role] = {
            "occurrence_count": occurrence,
            "node_count": occurrence,
            "occurrence_ids": [f"{role}:{node_ids[role]}"] if occurrence else [],
            "membership_counts": {
                "none": 0,
                "anchor-only": occurrence if membership == "anchor-only" else 0,
                "complete": occurrence if membership == "complete" else 0,
            },
            "minimum": policy["minimum"],
            "maximum": policy["maximum"],
            "style": policy["style"],
            "output": policy["output"],
        }
    return result


def test_frozen_audit() -> None:
    profile = load_profile("lecture-notes-en-zh")
    with tempfile.TemporaryDirectory(prefix="v23-docx-test-") as temporary:
        work = Path(temporary)
        output = work / "output"
        output.mkdir()
        profile_path = work / "profile.json"
        write_json(profile_path, profile)
        inventory = make_inventory(profile)
        theorem_source = "Theorem 1. Frozen source statement."
        definition_source = "Definition 2. Scoped anchor source statement."
        equation_source = "E equals m c squared."
        ir = {
            "schema_version": 2,
            "profile": {"id": profile["id"], "sha256": canonical_profile_sha256(profile)},
            "nodes": [
                {
                    "id": "theorem-node",
                    "source": {"text": theorem_source, "sha256": "test"},
                    "semantic": {"role": "theorem", "style": "theorem", "output": "bilingual"},
                },
                {
                    "id": "definition-node",
                    "source": {"text": definition_source, "sha256": "test"},
                    "semantic": {
                        "role": "definition",
                        "style": "definition",
                        "output": "bilingual",
                    },
                },
                {
                    "id": "equation-node",
                    "source": {"text": equation_source, "sha256": "test"},
                    "semantic": {"role": "equation", "style": "equation", "output": "source-only"},
                },
            ],
            "semantic_groups": [
                {
                    "id": "theorem:theorem-node",
                    "role": "theorem",
                    "identifier": None,
                    "anchor_node_id": "theorem-node",
                    "member_node_ids": ["theorem-node"],
                    "membership": "complete",
                },
                {
                    "id": "equation:equation-node",
                    "role": "equation",
                    "identifier": None,
                    "anchor_node_id": "equation-node",
                    "member_node_ids": ["equation-node"],
                    "membership": "anchor-only",
                },
                {
                    "id": "definition:definition-node",
                    "role": "definition",
                    "identifier": None,
                    "anchor_node_id": "definition-node",
                    "member_node_ids": ["definition-node"],
                    "membership": "anchor-only",
                },
            ],
            "inventories": {"role_inventory": inventory},
        }
        ir_path = work / "document-ir.json"
        write_json(ir_path, ir)
        markdown_path = output / "fixture.md"
        markdown_path.write_text(
            "\n".join(
                [
                    "<!-- bilingual:segment id=theorem-node source_sha256=test -->",
                    theorem_source,
                    "",
                    "> 定理一。这是冻结的目标陈述。",
                    "",
                    "<!-- bilingual:segment id=definition-node source_sha256=test -->",
                    definition_source,
                    "",
                    "> 定义二。这是只包含锚点的目标陈述。",
                    "",
                    "<!-- bilingual:source-only id=equation-node source_sha256=test -->",
                    equation_source,
                    "",
                ]
            ),
            encoding="utf-8",
        )
        build = {
            "schema_version": 1,
            "profile_id": profile["id"],
            "profile_file_sha256": sha256_file(profile_path),
            "document_ir_sha256": sha256_file(ir_path),
            "role_inventory": inventory,
            "external_uris": [],
            "assets": [],
            "markdown": markdown_path.name,
            "markdown_sha256": sha256_file(markdown_path),
        }
        write_json(output / "build-manifest.json", build)

        docx_path = output / "fixture.docx"
        if shutil.which("pandoc"):
            build_process = subprocess.run(
                [
                    sys.executable,
                    str(SCRIPT_DIR / "build_docx.py"),
                    str(markdown_path),
                    str(docx_path),
                    "--work-dir",
                    str(work),
                    "--expected-role",
                    "theorem=1",
                    "--expected-role",
                    "definition=1",
                    "--expected-role",
                    "equation=1",
                    "--expected-problems",
                    "0",
                ],
                check=False,
                capture_output=True,
                text=True,
            )
            assert build_process.returncode == 0, build_process.stdout + build_process.stderr
        else:
            document = Document()
            document.styles.add_style("Block Text", WD_STYLE_TYPE.PARAGRAPH)
            add_generic_range(
                document,
                "theorem",
                "theorem",
                "theorem:theorem-node",
                theorem_source,
                "定理一。这是冻结的目标陈述。",
            )
            add_generic_range(
                document,
                "definition",
                "definition",
                "definition:definition-node",
                definition_source,
                "定义二。这是只包含锚点的目标陈述。",
                membership="anchor-only",
            )
            document.add_paragraph(equation_source)
            docx_style.configure_profile(profile)
            docx_style.apply_styles(
                document,
                document_title="V2.3 frozen audit",
                header_label="test",
                footer_label="测试",
            )
            document.save(docx_path)
        audit_path = output / "docx-audit.json"
        process = subprocess.run(
            [
                sys.executable,
                str(SCRIPT_DIR / "audit_docx.py"),
                str(docx_path),
                "--work-dir",
                str(work),
                "--expected-role",
                "theorem=1",
                "--expected-role",
                "definition=1",
                "--expected-role",
                "equation=1",
                "--output",
                str(audit_path),
            ],
            check=False,
            capture_output=True,
            text=True,
        )
        assert process.returncode == 0, process.stdout + process.stderr
        report = json.loads(audit_path.read_text(encoding="utf-8"))
        assert report["status"] == "passed"
        assert report["role_counts"]["theorem"] == 1
        assert report["role_counts"]["definition"] == 1
        assert report["scoped_anchor_callout_checks"] == {
            "definition:definition-node": True
        }
        assert report["non_structural_anchor_unboxed"] == {
            "equation:equation-node": True
        }
        assert report["source_only_occurrence_counts"] == {"equation-node": 1}
        assert report["document_ir_sha256"] == sha256_file(ir_path)

        def audit_fixture(path: Path) -> subprocess.CompletedProcess[str]:
            return subprocess.run(
                [
                    sys.executable,
                    str(SCRIPT_DIR / "audit_docx.py"),
                    str(path),
                    "--work-dir",
                    str(work),
                    "--expected-role",
                    "theorem=1",
                    "--expected-role",
                    "definition=1",
                    "--expected-role",
                    "equation=1",
                ],
                check=False,
                capture_output=True,
                text=True,
            )

        missing_document = Document()
        missing_document.styles.add_style("Block Text", WD_STYLE_TYPE.PARAGRAPH)
        add_generic_range(
            missing_document,
            "theorem",
            "theorem",
            "theorem:theorem-node",
            theorem_source,
            "定理一。这是冻结的目标陈述。",
        )
        missing_document.add_paragraph(definition_source)
        missing_document.add_paragraph("定义二。这是只包含锚点的目标陈述。")
        missing_document.add_paragraph(equation_source)
        docx_style.configure_profile(profile)
        docx_style.apply_styles(
            missing_document,
            document_title="missing scoped anchor",
            header_label="test",
            footer_label="测试",
        )
        missing_path = output / "missing-anchor-callout.docx"
        missing_document.save(missing_path)
        missing_audit = audit_fixture(missing_path)
        assert missing_audit.returncode == 1
        assert "scoped_anchor_callouts_are_structurally_stable" in missing_audit.stdout

        absorbed_document = Document()
        absorbed_document.styles.add_style("Block Text", WD_STYLE_TYPE.PARAGRAPH)
        add_generic_range(
            absorbed_document,
            "theorem",
            "theorem",
            "theorem:theorem-node",
            theorem_source,
            "定理一。这是冻结的目标陈述。",
        )
        absorbed_document.add_paragraph(
            f"{GENERIC_BEGIN}|definition|definition|anchor-only|definition:definition-node",
            style="Block Text",
        )
        absorbed_document.add_paragraph(definition_source, style="Block Text")
        absorbed_document.add_paragraph(equation_source, style="Block Text")
        absorbed_document.add_paragraph("", style="Block Text")
        absorbed_document.add_paragraph(
            "定义二。这是只包含锚点的目标陈述。", style="Block Text"
        )
        absorbed_document.add_paragraph(
            f"{GENERIC_END}|definition|definition|anchor-only|definition:definition-node",
            style="Block Text",
        )
        docx_style.configure_profile(profile)
        docx_style.apply_styles(
            absorbed_document,
            document_title="absorbed neighbor",
            header_label="test",
            footer_label="测试",
        )
        absorbed_path = output / "absorbed-neighbor.docx"
        absorbed_document.save(absorbed_path)
        absorbed_audit = audit_fixture(absorbed_path)
        assert absorbed_audit.returncode == 1
        assert "scoped_anchor_callouts_are_structurally_stable" in absorbed_audit.stdout
        assert "non_structural_anchor_groups_are_not_boxed" in absorbed_audit.stdout

        compile_context = compile_docx_pdf.load_v2_context(work, docx_path)
        assert compile_context["schema_version"] == 2
        assert compile_context["build"]["role_inventory"] == inventory
        assert compile_context["docx_audit"]["status"] == "passed"
        assert compile_context["docx_audit_bindings"]["docx_sha256"] == sha256_file(
            docx_path
        )
        assert compile_docx_pdf.parse_expected_roles(["theorem=1", "equation=1"]) == {
            "theorem": 1,
            "equation": 1,
        }

        audit_bytes = audit_path.read_bytes()
        audit_path.unlink()
        try:
            compile_docx_pdf.load_v2_context(work, docx_path)
        except ValueError as exc:
            assert "missing DOCX audit" in str(exc)
        else:
            raise AssertionError("schema V2 compile accepted a missing DOCX audit")
        audit_path.write_bytes(audit_bytes)

        docx_bytes = docx_path.read_bytes()
        docx_path.write_bytes(docx_bytes + b"changed after DOCX audit")
        try:
            compile_docx_pdf.load_v2_context(work, docx_path)
        except ValueError as exc:
            assert "docx_sha256" in str(exc)
        else:
            raise AssertionError("schema V2 compile accepted modified DOCX bytes")
        docx_path.write_bytes(docx_bytes)

        stale_audit = json.loads(audit_bytes.decode("utf-8"))
        stale_audit["profile_file_sha256"] = "0" * 64
        write_json(audit_path, stale_audit)
        try:
            compile_docx_pdf.load_v2_context(work, docx_path)
        except ValueError as exc:
            assert "profile_file_sha256" in str(exc)
        else:
            raise AssertionError("schema V2 compile accepted stale audit bindings")
        audit_path.write_bytes(audit_bytes)


def main() -> None:
    test_legacy_transform()
    test_generic_ast()
    test_shared_style_roles()
    test_html_table_materialization()
    test_frozen_audit()
    print(
        "V2.3 DOCX tests passed: legacy, structural AST, shared styles, "
        "native tables, frozen audit"
    )


if __name__ == "__main__":
    main()
