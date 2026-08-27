#!/usr/bin/env python3
from __future__ import annotations

import hashlib
import io
import json
import os
import subprocess
import sys
import tempfile
import unittest
import zipfile
from contextlib import redirect_stdout
from pathlib import Path
from unittest.mock import patch

from PIL import Image


REPOSITORY = Path(__file__).resolve().parents[1]
SCRIPTS = REPOSITORY / "skills" / "make-bilingual-study-pdf" / "scripts"
sys.dont_write_bytecode = True
sys.path.insert(0, str(SCRIPTS))

from docx import Document  # noqa: E402
import build_docx as build_docx_module  # noqa: E402
from audit_source import current_source_audit_bindings  # noqa: E402
from document_ir import expected_ir  # noqa: E402
from profile import canonical_profile_sha256, load_profile  # noqa: E402


class DocxArtifactBoundaryTests(unittest.TestCase):
    @staticmethod
    def digest(payload: bytes) -> str:
        return hashlib.sha256(payload).hexdigest()

    @staticmethod
    def write_json(path: Path, value: object) -> None:
        path.write_text(
            json.dumps(value, ensure_ascii=False, indent=2) + "\n",
            encoding="utf-8",
        )

    def run_script(
        self, name: str, *arguments: Path | str
    ) -> subprocess.CompletedProcess[str]:
        environment = dict(os.environ)
        environment["PYTHONDONTWRITEBYTECODE"] = "1"
        return subprocess.run(
            [sys.executable, str(SCRIPTS / name), *(str(item) for item in arguments)],
            check=False,
            capture_output=True,
            text=True,
            timeout=20,
            env=environment,
        )

    def make_work(self, root: Path, *, schema_v2: bool = False) -> Path:
        work = root / "work"
        work.mkdir()
        profile = load_profile(
            "lecture-notes-en-zh" if schema_v2 else "assignment-en-zh"
        )
        (work / "profile.json").write_text(
            json.dumps(profile, ensure_ascii=False, indent=2) + "\n",
            encoding="utf-8",
        )
        return work

    def make_v1_output_gate(self, work: Path) -> tuple[Path, Path]:
        output = work / "output"
        output.mkdir(exist_ok=True)
        profile = load_profile("assignment-en-zh")
        self.write_json(work / "document-ir.json", {"fixture": True})
        (work / "blocks.jsonl").write_bytes(b"")
        (work / "oracle.txt").write_text("fixture\f", encoding="utf-8")
        (work / "oracle-layout.txt").write_text(
            "fixture layout\f", encoding="utf-8"
        )
        renders = work / "renders"
        renders.mkdir(exist_ok=True)
        Image.new("RGB", (2, 3), "white").save(renders / "page-1.png")
        source_contact = work / "source-contact"
        source_contact.mkdir(exist_ok=True)
        source_contact_path = source_contact / "contact-1.png"
        Image.new("RGB", (2, 3), "white").save(source_contact_path)
        source_pdf = work / "source.pdf"
        source_pdf.write_bytes(b"source PDF fixture\n")
        self.write_json(
            work / "manifest.json",
            {
                "source_pdf": str(source_pdf.absolute()),
                "source_sha256": self.digest(source_pdf.read_bytes()),
                "page_count": 1,
                "artifacts": {
                    "profile": "profile.json",
                    "document_ir": "document-ir.json",
                    "blocks": "blocks.jsonl",
                    "oracle": "oracle.txt",
                    "oracle_layout": "oracle-layout.txt",
                    "renders": "renders/page-*.png",
                    "source_contact": "source-contact/contact-*.png",
                },
                "visuals": [],
                "source_contact_sheets": [
                    {
                        "path": "source-contact/contact-1.png",
                        "sha256": self.digest(source_contact_path.read_bytes()),
                        "first_page": 1,
                        "last_page": 1,
                    }
                ],
            },
        )
        self.write_json(work / "document-ir.json", expected_ir(work, profile))
        self.write_json(
            work / "source-audit.json",
            {
                "status": "passed",
                **current_source_audit_bindings(work),
                "problem_ids": {
                    "oracle": [],
                    "extracted": [],
                    "missing": [],
                    "extra": [],
                },
                "global_fivegram_coverage": 1.0,
                "minimum_global_coverage": profile["qa"][
                    "minimum_global_fivegram_coverage"
                ],
                "rendered_pages": 1,
                "page_results": [
                    {
                        "page": 1,
                        "matched_fivegrams": 0,
                        "oracle_fivegrams": 0,
                        "block_count": 0,
                        "coverage": 1.0,
                    }
                ],
                "failures": [],
            },
        )
        translation = work / "translation"
        translation.mkdir(exist_ok=True)
        (translation / "requests").mkdir(exist_ok=True)
        (translation / "responses").mkdir(exist_ok=True)
        glossary = translation / "glossary.json"
        self.write_json(glossary, {"terms": []})
        plan = translation / "plan.json"
        self.write_json(
            plan,
            {
                "schema_version": 2,
                "profile_id": profile["id"],
                "profile_sha256": canonical_profile_sha256(profile),
                "profile_file_sha256": self.digest((work / "profile.json").read_bytes()),
                "document_ir_sha256": self.digest(
                    (work / "document-ir.json").read_bytes()
                ),
                "source_pdf_sha256": self.digest(source_pdf.read_bytes()),
                "source_manifest_sha256": self.digest(
                    (work / "manifest.json").read_bytes()
                ),
                "source_blocks_sha256": self.digest((work / "blocks.jsonl").read_bytes()),
                "source_audit_sha256": self.digest(
                    (work / "source-audit.json").read_bytes()
                ),
                "glossary_sha256": self.digest(glossary.read_bytes()),
                "target_language": profile["translation"]["target_language"],
                "batch_count": 0,
                "batches": [],
                "expected_segment_count": 0,
                "expected_ids": [],
            },
        )
        merged = translation / "translations-merged.jsonl"
        merged.write_bytes(b"")
        translation_audit = translation / "translation-audit.json"
        self.write_json(
            translation_audit,
            {
                "status": "passed",
                "plan_sha256": self.digest(plan.read_bytes()),
                "merged_sha256": self.digest(merged.read_bytes()),
                "response_bindings": [],
                "response_files": [],
                "merged_output": "translation/translations-merged.jsonl",
                "expected_segments": 0,
                "response_segments": 0,
                "validated_segments": 0,
                "missing_ids": [],
                "extra_ids": [],
                "duplicate_ids": [],
                "invalid_source_hash_ids": [],
                "empty_translation_ids": [],
                "untranslated_ids": [],
                "source_copy_ids": [],
                "placeholder_failures": {},
                "glossary_failures": {},
                "failures": [],
            },
        )
        markdown = output / "fixture.md"
        markdown.write_text("English source.\n\n中文目标。\n", encoding="utf-8")
        build = output / "build-manifest.json"
        self.write_json(
            build,
            {
                "markdown": markdown.name,
                "markdown_sha256": self.digest(markdown.read_bytes()),
                "latex": None,
                "assets": [],
                "block_count": 0,
                "disposition_counts": {},
                "role_inventory": {
                    role: {"occurrence_count": 0}
                    for role in ("problem", "example", "tip")
                },
                "problem_ids": [],
                "external_uris": [],
                "translation_audit_sha256": self.digest(
                    translation_audit.read_bytes()
                ),
            },
        )
        self.write_json(
            output / "output-audit.json",
            {
                "status": "passed",
                "build_manifest_sha256": self.digest(build.read_bytes()),
                "artifact_bindings": {
                    "markdown": {
                        "path": markdown.name,
                        "sha256": self.digest(markdown.read_bytes()),
                    },
                    "assets": [],
                },
                "markdown": markdown.name,
                "latex": None,
                "asset_count": 0,
                "block_count": 0,
                "disposition_counts": {},
                "semantic_constraint_checks": {"fixture": True},
                "failures": [],
            },
        )
        return output, markdown

    @staticmethod
    def make_docx(path: Path) -> None:
        document = Document()
        document.add_paragraph("English source paragraph.")
        document.add_paragraph("中文目标段落。")
        document.save(path)

    @staticmethod
    def snapshot_tree(root: Path) -> tuple[tuple[str, str, bytes | None], ...]:
        records = []
        for path in sorted(root.rglob("*"), key=lambda item: item.as_posix()):
            relative = path.relative_to(root).as_posix()
            if path.is_dir():
                records.append((relative, "directory", None))
            else:
                records.append((relative, "file", path.read_bytes()))
        return tuple(records)

    def test_work_mode_ast_and_style_happy_paths_are_atomic(self) -> None:
        with tempfile.TemporaryDirectory(
            prefix="docx-boundary-happy-"
        ) as temporary:
            root = Path(temporary)
            work = self.make_work(root)
            stage = work / "output" / "docx-build"
            stage.mkdir(parents=True)
            ast_input = stage / "source.json"
            ast_output = stage / "grouped.json"
            ast_input.write_text(
                json.dumps(
                    {"pandoc-api-version": [1, 23], "meta": {}, "blocks": []}
                ),
                encoding="utf-8",
            )
            transformed = self.run_script(
                "docx_ast.py", ast_input, ast_output, "--work-dir", work
            )
            self.assertEqual(
                transformed.returncode, 0, transformed.stdout + transformed.stderr
            )
            self.assertIsInstance(json.loads(ast_output.read_text(encoding="utf-8")), dict)

            raw_docx = stage / "raw.docx"
            styled_docx = stage / "styled.docx"
            document = Document()
            document.add_paragraph("English source paragraph.")
            document.add_paragraph("中文目标段落。")
            document.save(raw_docx)
            styled = self.run_script(
                "docx_style.py", raw_docx, styled_docx, "--work-dir", work
            )
            self.assertEqual(styled.returncode, 0, styled.stdout + styled.stderr)
            with zipfile.ZipFile(styled_docx) as archive:
                self.assertIn("word/document.xml", archive.namelist())
            self.assertFalse(list(stage.glob(".*.tmp")))

    def test_work_mode_hardlink_targets_are_rejected_without_external_writes(self) -> None:
        with tempfile.TemporaryDirectory(
            prefix="docx-boundary-hardlink-"
        ) as temporary:
            root = Path(temporary)
            work = self.make_work(root)
            stage = work / "output" / "docx-build"
            stage.mkdir(parents=True)
            ast_input = stage / "source.json"
            ast_input.write_text(
                '{"pandoc-api-version":[1,23],"meta":{},"blocks":[]}',
                encoding="utf-8",
            )
            outside = root / "outside.bin"
            outside.write_bytes(b"outside sentinel")
            target = stage / "grouped.json"
            os.link(outside, target)
            completed = self.run_script(
                "docx_ast.py", ast_input, target, "--work-dir", work
            )
            self.assertNotEqual(completed.returncode, 0)
            self.assertIn("hard-linked", completed.stdout + completed.stderr)
            self.assertEqual(outside.read_bytes(), b"outside sentinel")

    def test_work_mode_ast_and_style_reject_role_aliases_before_mutation(self) -> None:
        with tempfile.TemporaryDirectory(prefix="docx-helper-role-alias-") as temporary:
            root = Path(temporary)
            work = self.make_work(root)
            stage = work / "output" / "docx-build"
            stage.mkdir(parents=True)
            ast_input = stage / "source.json"
            ast_input.write_text(
                '{"pandoc-api-version":[1,23],"meta":{},"blocks":[]}',
                encoding="utf-8",
            )
            manifest = work / "manifest.json"
            manifest.write_bytes(b"manifest sentinel\n")
            raw_docx = stage / "raw.docx"
            self.make_docx(raw_docx)
            before = self.snapshot_tree(work)

            ast = self.run_script(
                "docx_ast.py", ast_input, manifest, "--work-dir", work
            )
            styled = self.run_script(
                "docx_style.py", raw_docx, raw_docx, "--work-dir", work
            )

            self.assertNotEqual(ast.returncode, 0)
            self.assertIn("unsafe path", ast.stdout + ast.stderr)
            self.assertNotEqual(styled.returncode, 0)
            self.assertIn("roles must be distinct", styled.stdout + styled.stderr)
            self.assertEqual(self.snapshot_tree(work), before)

        for name in ("fixture.docx", "docx-audit.json"):
            with self.subTest(name=name), tempfile.TemporaryDirectory(
                prefix=f"docx-fixed-{name}-"
            ) as temporary:
                root = Path(temporary)
                work = self.make_work(root, schema_v2=True)
                output = work / "output"
                output.mkdir()
                source = output / "fixture.md"
                source.write_text("source", encoding="utf-8")
                outside = root / "outside.bin"
                outside.write_bytes(b"outside sentinel")
                os.link(outside, output / name)
                completed = self.run_script(
                    "build_docx.py",
                    source,
                    output / "fixture.docx",
                    "--work-dir",
                    work,
                )
                self.assertNotEqual(completed.returncode, 0)
                self.assertIn("hard-linked", completed.stdout + completed.stderr)
                self.assertEqual(outside.read_bytes(), b"outside sentinel")

    def test_docx_build_junction_is_rejected_without_external_writes(self) -> None:
        with tempfile.TemporaryDirectory(
            prefix="docx-build-junction-"
        ) as temporary:
            root = Path(temporary)
            work = self.make_work(root, schema_v2=True)
            output = work / "output"
            output.mkdir()
            source = output / "fixture.md"
            source.write_text("source", encoding="utf-8")
            outside = root / "outside"
            outside.mkdir()
            sentinel = outside / "sentinel.bin"
            sentinel.write_bytes(b"outside sentinel")
            build_dir = output / "docx-build"
            if os.name == "nt":
                created = subprocess.run(
                    ["cmd.exe", "/d", "/c", "mklink", "/J", str(build_dir), str(outside)],
                    check=False,
                    capture_output=True,
                    text=True,
                    timeout=5,
                )
                if created.returncode != 0:
                    self.fail(
                        "Windows junction regression must execute: "
                        f"{created.stdout}{created.stderr}"
                    )
            else:
                os.symlink(outside, build_dir, target_is_directory=True)
            try:
                completed = self.run_script(
                    "build_docx.py",
                    source,
                    output / "fixture.docx",
                    "--work-dir",
                    work,
                )
                self.assertNotEqual(completed.returncode, 0)
                marker = "reparse points" if os.name == "nt" else "symbolic links"
                self.assertIn(marker, completed.stdout + completed.stderr)
                self.assertEqual(sentinel.read_bytes(), b"outside sentinel")
            finally:
                if os.path.lexists(build_dir):
                    if os.name == "nt":
                        os.rmdir(build_dir)
                    else:
                        build_dir.unlink()

    def test_v1_work_audit_binds_docx_and_output_freeze_atomically(self) -> None:
        with tempfile.TemporaryDirectory(prefix="docx-v1-audit-") as temporary:
            root = Path(temporary)
            work = self.make_work(root)
            output, _markdown = self.make_v1_output_gate(work)
            docx = output / "fixture.docx"
            self.make_docx(docx)
            completed = self.run_script(
                "audit_docx.py",
                docx,
                "--work-dir",
                work,
                "--profile",
                work / "profile.json",
                "--expected-problems",
                "0",
                "--expected-examples",
                "0",
                "--expected-tips",
                "0",
                "--expected-links",
                "0",
                "--output",
                output / "docx-audit.json",
            )
            self.assertEqual(completed.returncode, 0, completed.stdout + completed.stderr)
            report = json.loads(
                (output / "docx-audit.json").read_text(encoding="utf-8")
            )
            self.assertEqual(report["status"], "passed")
            self.assertEqual(report["failures"], [])
            self.assertEqual(report["docx_sha256"], self.digest(docx.read_bytes()))
            self.assertEqual(
                report["output_audit_sha256"],
                self.digest((output / "output-audit.json").read_bytes()),
            )

    def test_v1_work_audit_rejects_hardlinked_input_and_report(self) -> None:
        for role in ("docx", "audit"):
            with self.subTest(role=role), tempfile.TemporaryDirectory(
                prefix=f"docx-v1-{role}-hardlink-"
            ) as temporary:
                root = Path(temporary)
                work = self.make_work(root)
                output, _markdown = self.make_v1_output_gate(work)
                docx = output / "fixture.docx"
                self.make_docx(docx)
                outside = root / "outside.bin"
                outside.write_bytes(b"outside sentinel")
                target = docx if role == "docx" else output / "docx-audit.json"
                if target.exists():
                    target.unlink()
                os.link(outside, target)
                completed = self.run_script(
                    "audit_docx.py",
                    docx,
                    "--work-dir",
                    work,
                    "--profile",
                    work / "profile.json",
                    "--output",
                    output / "docx-audit.json",
                )
                self.assertNotEqual(completed.returncode, 0)
                self.assertIn("hard-linked", completed.stdout + completed.stderr)
                self.assertEqual(outside.read_bytes(), b"outside sentinel")

    def test_v1_work_build_uses_temporary_stage_without_clobbering_work_root(self) -> None:
        with tempfile.TemporaryDirectory(prefix="docx-v1-build-") as temporary:
            root = Path(temporary)
            work = self.make_work(root)
            output, markdown = self.make_v1_output_gate(work)
            sentinels: dict[Path, bytes] = {}
            for name in ("source.json", "grouped.json", "raw.docx", "styled.docx"):
                path = work / name
                path.write_bytes(f"root sentinel {name}".encode("utf-8"))
                sentinels[path] = path.read_bytes()

            def fake_pandoc(command: list[str]) -> None:
                destination = Path(command[command.index("--output") + 1])
                if command[command.index("--to") + 1] == "json":
                    destination.write_text(
                        json.dumps(
                            {"pandoc-api-version": [1, 23], "meta": {}, "blocks": []}
                        ),
                        encoding="utf-8",
                    )
                else:
                    self.make_docx(destination)

            stdout = io.StringIO()
            with (
                patch.object(
                    sys,
                    "argv",
                    [
                        "build_docx.py",
                        str(markdown),
                        str(output / "fixture.docx"),
                        "--resource-path",
                        str(output),
                        "--profile",
                        str(work / "profile.json"),
                        "--expected-problems",
                        "0",
                        "--work-dir",
                        str(work),
                    ],
                ),
                patch.object(build_docx_module, "run", side_effect=fake_pandoc),
                patch.object(build_docx_module.shutil, "which", return_value="pandoc"),
                redirect_stdout(stdout),
            ):
                build_docx_module.main()
            self.assertTrue((output / "fixture.docx").is_file())
            for path, payload in sentinels.items():
                self.assertEqual(path.read_bytes(), payload)
            self.assertFalse(list(output.glob(".docx-build-*")))

    def test_v1_build_rejects_staging_input_alias_before_tool_discovery(self) -> None:
        with tempfile.TemporaryDirectory(prefix="docx-v1-stage-alias-") as temporary:
            root = Path(temporary)
            work = self.make_work(root)
            output, _markdown = self.make_v1_output_gate(work)
            stage = work / "scratch"
            stage.mkdir()
            source = stage / "source.json"
            source.write_bytes(b"source sentinel")
            completed = self.run_script(
                "build_docx.py",
                source,
                output / "fixture.docx",
                "--resource-path",
                stage,
                "--profile",
                work / "profile.json",
                "--work-dir",
                work,
                "--build-dir",
                stage,
            )
            self.assertNotEqual(completed.returncode, 0)
            self.assertIn("output/docx-build", completed.stdout + completed.stderr)
            self.assertEqual(source.read_bytes(), b"source sentinel")

    def test_work_build_rejects_reference_and_profile_role_aliases_before_pandoc(self) -> None:
        for role in ("reference-output", "reference-gate", "profile-output"):
            with self.subTest(role=role), tempfile.TemporaryDirectory(
                prefix=f"docx-build-role-alias-{role}-"
            ) as temporary:
                root = Path(temporary)
                work = self.make_work(root)
                output, markdown = self.make_v1_output_gate(work)
                output_docx = output / "fixture.docx"
                arguments = [
                    "build_docx.py",
                    str(markdown),
                    str(output_docx),
                    "--resource-path",
                    str(output),
                    "--work-dir",
                    str(work),
                ]
                if role == "reference-output":
                    output_docx.write_bytes(b"reference template sentinel\n")
                    arguments.extend(("--profile", str(work / "profile.json")))
                    arguments.extend(("--reference-doc", str(output_docx)))
                elif role == "reference-gate":
                    arguments.extend(("--profile", str(work / "profile.json")))
                    arguments.extend(
                        ("--reference-doc", str(output / "output-audit.json"))
                    )
                else:
                    output_docx.write_bytes((work / "profile.json").read_bytes())
                    arguments.extend(("--profile", str(output_docx)))
                before = self.snapshot_tree(work)

                with (
                    patch.object(sys, "argv", arguments),
                    patch.object(
                        build_docx_module.shutil,
                        "which",
                        side_effect=AssertionError("Pandoc discovery must not run"),
                    ),
                    self.assertRaises(SystemExit) as caught,
                ):
                    build_docx_module.main()

                self.assertTrue(
                    "alias" in str(caught.exception)
                    or "canonical WORK/profile.json" in str(caught.exception),
                    str(caught.exception),
                )
                self.assertEqual(self.snapshot_tree(work), before)

    def test_work_build_rejects_unfrozen_markdown_resource_and_reference_before_pandoc(
        self,
    ) -> None:
        for role in (
            "rogue-markdown",
            "external-resource",
            "external-reference",
            "work-reference",
        ):
            with self.subTest(role=role), tempfile.TemporaryDirectory(
                prefix=f"docx-frozen-role-{role}-"
            ) as temporary:
                root = Path(temporary)
                work = self.make_work(root)
                output, markdown = self.make_v1_output_gate(work)
                source = markdown
                resource = output
                reference: Path | None = None
                outside = root / "outside"
                outside.mkdir()
                outside_sentinel = outside / "sentinel.bin"
                outside_sentinel.write_bytes(b"outside sentinel")
                if role == "rogue-markdown":
                    source = output / "rogue.md"
                    source.write_text("rogue source\n", encoding="utf-8")
                elif role == "external-resource":
                    resource = outside
                elif role == "external-reference":
                    reference = outside / "reference.docx"
                    reference.write_bytes(b"external reference sentinel")
                else:
                    reference = output / "reference.docx"
                    reference.write_bytes(b"WORK reference sentinel")
                arguments = [
                    "build_docx.py",
                    str(source),
                    str(output / "fixture.docx"),
                    "--resource-path",
                    str(resource),
                    "--profile",
                    str(work / "profile.json"),
                    "--work-dir",
                    str(work),
                ]
                if reference is not None:
                    arguments.extend(("--reference-doc", str(reference)))
                before = self.snapshot_tree(work)
                outside_before = self.snapshot_tree(outside)
                with (
                    patch.object(sys, "argv", arguments),
                    patch.object(
                        build_docx_module.shutil,
                        "which",
                        side_effect=AssertionError("Pandoc discovery must not run"),
                    ),
                    self.assertRaises(SystemExit),
                ):
                    build_docx_module.main()
                self.assertEqual(self.snapshot_tree(work), before)
                self.assertEqual(self.snapshot_tree(outside), outside_before)


if __name__ == "__main__":
    unittest.main(verbosity=2)
