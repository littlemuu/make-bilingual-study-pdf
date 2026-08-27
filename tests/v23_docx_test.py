#!/usr/bin/env python3
"""Focused V2.3 DOCX semantic-container and frozen-audit regression tests."""
from __future__ import annotations

import json
import os
import shutil
import subprocess
import sys
import tempfile
from pathlib import Path

import pymupdf as fitz
from PIL import Image

REPOSITORY = Path(__file__).resolve().parents[1]
SCRIPTS = REPOSITORY / "skills" / "make-bilingual-study-pdf" / "scripts"
sys.path.insert(0, str(SCRIPTS))

from docx import Document
from docx.enum.style import WD_STYLE_TYPE

import docx_style
import compile_docx_pdf
from adapters.mineru import LARGE_RASTER_PAGE_AREA_RATIO, RASTER_COVERAGE_METHOD
from audit_source import current_source_audit_bindings
from build_docx import materialize_html_tables
from common import (
    read_json,
    read_jsonl,
    sha256_file,
    sha256_text,
    write_json,
    write_jsonl,
)
from document_ir import expected_ir
from docx_ast import GENERIC_BEGIN, GENERIC_END, PROBLEM_BEGIN, PROBLEM_END, transform
from profile import canonical_profile_sha256, load_profile, profile_contract


SCRIPT_DIR = SCRIPTS


def paragraph(text: str) -> dict:
    return {"t": "Para", "c": [{"t": "Str", "c": text}]}


def marker(node_id: str) -> dict:
    return {
        "t": "RawBlock",
        "c": ["html", f"<!-- bilingual:segment id={node_id} source_sha256=test -->"],
    }


def target(text: str) -> dict:
    return {"t": "BlockQuote", "c": [paragraph(text)]}


def test_legacy_transform() -> None:
    profile = load_profile("assignment-en-zh")
    document = {
        "pandoc-api-version": [1, 23],
        "meta": {},
        "blocks": [
            {
                "t": "BlockQuote",
                "c": [paragraph("Problem (legacy) source"), paragraph("问题（legacy）目标")],
            }
        ],
    }
    result = transform(document, profile)
    grouped = result["blocks"][0]["c"]
    assert grouped[0] == paragraph(PROBLEM_BEGIN)
    assert grouped[-1] == paragraph(PROBLEM_END)
    assert sum(block.get("t") == "HorizontalRule" for block in grouped) == 1
    assert result["meta"]["v2-problem-group-count"]["c"] == "1"


def test_generic_ast() -> None:
    profile = load_profile("lecture-notes-en-zh")
    document = {
        "pandoc-api-version": [1, 23],
        "meta": {},
        "blocks": [
            marker("a"),
            paragraph("Theorem 1. Source anchor"),
            target("定理 1。目标锚点"),
            marker("b"),
            paragraph("Second source member"),
            target("第二个目标成员"),
            marker("c"),
            paragraph("Lemma 2. Anchor only"),
            target("引理 2。仅锚点"),
            marker("d"),
            paragraph("Ordinary neighboring paragraph"),
            target("普通相邻段落"),
        ],
    }
    groups = [
        {
            "id": "theorem:a",
            "role": "theorem",
            "member_node_ids": ["a", "b"],
            "membership": "complete",
        },
        {
            "id": "lemma:c",
            "role": "lemma",
            "anchor_node_id": "c",
            "member_node_ids": ["c"],
            "membership": "anchor-only",
        },
        {
            "id": "paragraph:d",
            "role": "paragraph",
            "anchor_node_id": "d",
            "member_node_ids": ["d"],
            "membership": "anchor-only",
        },
    ]
    result = transform(document, profile, semantic_groups=groups)
    callouts = [block for block in result["blocks"] if block.get("t") == "BlockQuote"]
    assert len(callouts) == 3
    structural = callouts[0]["c"]
    assert GENERIC_BEGIN in structural[0]["c"][0]["c"]
    assert GENERIC_END in structural[-1]["c"][0]["c"]
    assert sum(block.get("t") == "HorizontalRule" for block in structural) == 1
    structural_text = " ".join(
        str(inline.get("c", ""))
        for block in structural
        for inline in block.get("c", [])
        if isinstance(inline, dict)
    )
    assert structural_text.index("Theorem 1. Source anchor") < structural_text.index("Second source member")
    assert structural_text.index("Second source member") < structural_text.index("定理 1。目标锚点")
    scoped_text = " ".join(
        str(inline.get("c", ""))
        for block in callouts[1]["c"]
        for inline in block.get("c", [])
        if isinstance(inline, dict)
    )
    assert "Lemma 2. Anchor only" in scoped_text
    assert "Ordinary neighboring paragraph" not in scoped_text
    assert any(block.get("t") == "RawBlock" for block in result["blocks"])
    counts = json.loads(result["meta"]["v23-structural-group-counts"]["c"])
    complete_counts = json.loads(
        result["meta"]["v23-complete-structural-group-counts"]["c"]
    )
    scoped_counts = json.loads(
        result["meta"]["v23-anchor-only-callout-counts"]["c"]
    )
    assert counts["theorem"] == 1 and counts["lemma"] == 1
    assert complete_counts["theorem"] == 1 and complete_counts["lemma"] == 0
    assert scoped_counts["theorem"] == 0 and scoped_counts["lemma"] == 1
    assert counts["paragraph"] == 0

    grouped_alias_document = {
        "pandoc-api-version": [1, 23],
        "meta": {},
        "blocks": [
            marker("b"),
            marker("a"),
            paragraph("Theorem 1. Source anchor Second source member"),
            target("定理一。目标锚点 第二个目标成员"),
        ],
    }
    grouped_alias = transform(
        grouped_alias_document, profile, semantic_groups=groups[:1]
    )
    assert len(grouped_alias["blocks"]) == 1
    assert grouped_alias["blocks"][0]["t"] == "BlockQuote"


def add_generic_range(
    document: Document,
    role: str,
    style: str,
    group_id: str,
    source: str,
    target_text: str,
    *,
    membership: str = "complete",
) -> None:
    document.add_paragraph(
        f"{GENERIC_BEGIN}|{role}|{style}|{membership}|{group_id}",
        style="Block Text",
    )
    document.add_paragraph(source, style="Block Text")
    document.add_paragraph("", style="Block Text")
    document.add_paragraph(target_text, style="Block Text")
    document.add_paragraph(
        f"{GENERIC_END}|{role}|{style}|{membership}|{group_id}",
        style="Block Text",
    )


def test_shared_style_roles() -> None:
    profile = load_profile("lecture-notes-en-zh")
    document = Document()
    document.styles.add_style("Block Text", WD_STYLE_TYPE.PARAGRAPH)
    add_generic_range(document, "theorem", "theorem", "theorem:a", "Theorem 1. A", "定理一")
    add_generic_range(
        document,
        "lemma",
        "theorem",
        "lemma:b",
        "Lemma 2. B",
        "引理二",
        membership="anchor-only",
    )
    document.add_paragraph("Section 1 is not a callout")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Metric"
    table.cell(0, 1).text = "Value"
    table.cell(1, 0).text = "Coverage"
    table.cell(1, 1).text = "100%"
    docx_style.configure_profile(profile)
    report = docx_style.apply_styles(
        document,
        document_title="V2.3 test",
        header_label="test",
        footer_label="测试",
    )
    assert report["role_callouts"]["theorem"] == 1
    assert report["role_callouts"]["lemma"] == 1
    assert report["complete_role_callouts"]["theorem"] == 1
    assert report["anchor_only_role_callouts"]["lemma"] == 1
    assert report["structural_occurrence_memberships"] == {
        "theorem:a": "complete",
        "lemma:b": "anchor-only",
    }
    assert report["structural_occurrence_ids"]["theorem"] == ["theorem:a"]
    assert report["structural_occurrence_ids"]["lemma"] == ["lemma:b"]
    assert report["table_count"] == 1
    assert report["table_cell_count"] == 4
    assert table.cell(0, 0)._tc.xpath("./w:tcPr/w:shd/@w:fill") == [
        docx_style.ACCENT_LIGHT
    ]
    assert table._tbl.xpath("./w:tblPr/w:tblBorders/w:insideV/@w:val") == [
        "single"
    ]
    assert all(GENERIC_BEGIN not in item.text and GENERIC_END not in item.text for item in document.paragraphs)


def test_html_table_materialization() -> None:
    if not shutil.which("pandoc"):
        return
    table_html = (
        "<table><thead><tr><th>Metric</th><th>Value</th></tr></thead>"
        "<tbody><tr><td>Coverage</td><td>100%</td></tr></tbody></table>"
    )
    parsed = subprocess.run(
        ["pandoc", "--from", "markdown", "--to", "json"],
        input=f"```{{=html}}\n{table_html}\n```\n",
        text=True,
        encoding="utf-8",
        capture_output=True,
        check=True,
    )
    ast = json.loads(parsed.stdout)
    assert ast["blocks"] == [{"t": "RawBlock", "c": ["html", table_html]}]
    materialized = materialize_html_tables(ast)
    assert [block["t"] for block in materialized["blocks"]] == ["Table"]

    reference_node = {
        "type": "list",
        "source": {
            "text": "- [1] Alpha (2025).\n- [2] Beta (2026).",
        },
    }
    assert "• [1] Alpha (2025). • [2] Beta (2026)." in (
        compile_docx_pdf.searchable_sources(reference_node)
    )


def run_work_script(name: str, work: Path, *arguments: str) -> None:
    environment = dict(os.environ)
    environment["PYTHONDONTWRITEBYTECODE"] = "1"
    environment["PYTHONUTF8"] = "1"
    completed = subprocess.run(
        [sys.executable, str(SCRIPT_DIR / name), str(work), *arguments],
        check=False,
        capture_output=True,
        text=True,
        encoding="utf-8",
        env=environment,
    )
    assert completed.returncode == 0, completed.stdout + completed.stderr


def prepare_frozen_v2_output(
    work: Path,
    profile: dict,
    theorem_source: str,
    definition_source: str,
    equation_source: str,
) -> tuple[dict, Path]:
    """Build a complete synthetic source-to-output freeze for the DOCX tests."""
    blocks = [
        {
            "id": "title-node",
            "page": 1,
            "bbox": [72.0, 30.0, 500.0, 50.0],
            "kind": "heading",
            "source": "Frozen Lecture Notes",
            "source_sha256": sha256_text("Frozen Lecture Notes"),
            "translatable": True,
            "links": [],
            "protected_spans": [],
            "stats": {},
            "adapter_role": "title",
        },
        {
            "id": "section-node",
            "page": 1,
            "bbox": [72.0, 55.0, 500.0, 75.0],
            "kind": "heading",
            "source": "1 Foundations",
            "source_sha256": sha256_text("1 Foundations"),
            "translatable": True,
            "links": [],
            "protected_spans": [],
            "stats": {},
            "adapter_role": "section",
        },
        {
            "id": "theorem-node",
            "page": 1,
            "bbox": [72.0, 80.0, 500.0, 110.0],
            "kind": "prose",
            "source": theorem_source,
            "source_sha256": sha256_text(theorem_source),
            "translatable": True,
            "links": [],
            "protected_spans": [],
            "stats": {},
            "adapter_role": "theorem",
        },
        {
            "id": "definition-node",
            "page": 1,
            "bbox": [72.0, 130.0, 500.0, 160.0],
            "kind": "prose",
            "source": definition_source,
            "source_sha256": sha256_text(definition_source),
            "translatable": True,
            "links": [],
            "protected_spans": [],
            "stats": {},
            "adapter_role": "definition",
        },
        {
            "id": "equation-node",
            "page": 1,
            "bbox": [72.0, 180.0, 500.0, 210.0],
            "kind": "equation",
            "source": equation_source,
            "source_sha256": sha256_text(equation_source),
            "translatable": False,
            "links": [],
            "protected_spans": [],
            "stats": {},
            "adapter_role": "equation",
        },
        {
            "id": "paragraph-node",
            "page": 1,
            "bbox": [72.0, 220.0, 500.0, 250.0],
            "kind": "prose",
            "source": "This paragraph closes the synthetic lecture fixture.",
            "source_sha256": sha256_text(
                "This paragraph closes the synthetic lecture fixture."
            ),
            "translatable": True,
            "links": [],
            "protected_spans": [],
            "stats": {},
            "adapter_role": "paragraph",
        },
    ]
    content: list[dict] = []
    evidence_items: list[dict] = []
    adapter_id = profile_contract(profile)["adapter"]
    for index, block in enumerate(blocks):
        pointer = f"/{index}"
        content_item = {
            "type": block["kind"],
            "sub_type": None,
            "text": block["source"],
            "page_idx": 0,
            "bbox": block["bbox"],
        }
        item_hash = sha256_text(
            json.dumps(
                content_item,
                ensure_ascii=False,
                sort_keys=True,
                separators=(",", ":"),
            )
        )
        block["evidence"] = {
            "adapter": adapter_id,
            "adapter_role": block["adapter_role"],
            "content_pointer": pointer,
            "content_item_pointer": pointer,
            "content_item_sha256": item_hash,
            "raw_type": block["kind"],
            "raw_sub_type": None,
        }
        if block["id"] == "theorem-node":
            block["evidence"]["structural_membership"] = {
                "status": "complete",
                "member_node_ids": ["theorem-node"],
            }
        content.append(content_item)
        evidence_items.append(
            {
                "pointer": pointer,
                "item_sha256": item_hash,
                "page_idx": 0,
                "page_order": index + 1,
                "raw_type": block["kind"],
                "raw_sub_type": None,
                "disposition": "emitted",
                "reason": None,
                "node_ids": [block["id"]],
                "visual_ids": [],
                "middle_pointers": [],
                "middle_match": "not-applicable",
            }
        )
    write_jsonl(work / "blocks.jsonl", blocks)

    (work / "oracle.txt").write_text("\f", encoding="utf-8")
    (work / "oracle-layout.txt").write_text("\f", encoding="utf-8")
    renders = work / "renders"
    renders.mkdir()
    Image.new("RGB", (2, 3), "white").save(renders / "page-1.png")
    source_contact = work / "source-contact"
    source_contact.mkdir()
    contact_path = source_contact / "contact-1.png"
    Image.new("RGB", (2, 3), "white").save(contact_path)

    source_pdf = work / "source.pdf"
    source_document = fitz.open()
    source_page = source_document.new_page(width=612, height=792)
    source_page.insert_text(
        (72, 72),
        "Lecture notes provide a deterministic native-text fixture for theorem, "
        "definition, and equation DOCX freeze-chain verification.",
    )
    source_pdf.write_bytes(
        source_document.tobytes(garbage=4, deflate=True, no_new_id=True)
    )
    source_document.close()

    adapter_inputs = work / "adapter-inputs"
    adapter_inputs.mkdir()
    origin_path = adapter_inputs / "fixture-origin.pdf"
    origin_path.write_bytes(source_pdf.read_bytes())
    content_path = adapter_inputs / "fixture-content.json"
    write_json(content_path, content)
    middle_path = adapter_inputs / "fixture-middle.json"
    write_json(middle_path, {})
    input_records = []
    for role, path in (
        ("origin", origin_path),
        ("content", content_path),
        ("middle", middle_path),
    ):
        input_records.append(
            {
                "role": role,
                "relative_path": path.name,
                "work_path": path.relative_to(work).as_posix(),
                "sha256": sha256_file(path),
                "size": path.stat().st_size,
            }
        )
    input_records.sort(key=lambda item: (item["role"], item["relative_path"]))
    evidence = {
        "schema_version": 1,
        "adapter": adapter_id,
        "source": {
            "logical_name": source_pdf.name,
            "sha256": sha256_file(source_pdf),
            "page_count": 1,
        },
        "mineru": {
            "version": "3.4.4",
            "backend": "pipeline",
            "support_level": "verified",
        },
        "raster_detection": {
            "method": RASTER_COVERAGE_METHOD,
            "large_page_area_ratio": LARGE_RASTER_PAGE_AREA_RATIO,
        },
        "inputs": input_records,
        "assets": [],
        "pages": [
            {
                "page_idx": 0,
                "page_size": [612.0, 792.0],
                "source_page_size": [612.0, 792.0],
                "native_text_characters": 130,
                "adapter_text_characters": sum(len(block["source"]) for block in blocks),
                "raster_image_area_ratio": 0.0,
                "manual_review_reasons": [],
                "status": "native_oracle_available",
            }
        ],
        "items": evidence_items,
        "manual_source_review_required": False,
        "manual_review_pages": [],
        "manual_review_page_comparisons": [],
        "manual_review_contact_sheets": [],
    }
    evidence_path = work / "adapter-evidence.json"
    write_json(evidence_path, evidence)
    manifest = {
        "schema_version": 4,
        "profile": {
            "id": profile["id"],
            "sha256": canonical_profile_sha256(profile),
        },
        "source_pdf": str(source_pdf),
        "source_sha256": sha256_file(source_pdf),
        "page_count": 1,
        "artifacts": {
            "profile": "profile.json",
            "document_ir": "document-ir.json",
            "adapter_evidence": "adapter-evidence.json",
            "blocks": "blocks.jsonl",
            "oracle": "oracle.txt",
            "oracle_layout": "oracle-layout.txt",
            "renders": "renders/page-*.png",
            "source_contact": "source-contact/contact-*.png",
        },
        "adapter": {
            "id": adapter_id,
            "evidence": "adapter-evidence.json",
            "evidence_sha256": sha256_file(evidence_path),
            "backend": "pipeline",
            "version": "3.4.4",
        },
        "input_artifacts": input_records,
        "external_uris": [],
        "external_uri_count": 0,
        "links": [],
        "visuals": [],
        "source_contact_sheets": [
            {
                "path": "source-contact/contact-1.png",
                "sha256": sha256_file(contact_path),
                "first_page": 1,
                "last_page": 1,
            }
        ],
        "source_review_pages": [],
        "source_review_contact_sheets": [],
    }
    write_json(work / "manifest.json", manifest)
    ir = expected_ir(work, profile)
    write_json(work / "document-ir.json", ir)
    source_bindings = current_source_audit_bindings(work)
    write_json(
        work / "source-audit.json",
        {
            "status": "passed",
            **source_bindings,
            "minimum_global_coverage": profile["qa"][
                "minimum_global_fivegram_coverage"
            ],
            "failures": [],
        },
    )

    translation_dir = work / "translation"
    translation_dir.mkdir()
    write_json(
        translation_dir / "glossary.json",
        {
            "schema_version": 1,
            "profile_id": profile["id"],
            "profile_sha256": canonical_profile_sha256(profile),
            "target_language": profile["translation"]["target_language"],
            "source_blocks_sha256": sha256_file(work / "blocks.jsonl"),
            "terms": [],
        },
    )
    run_work_script("prepare_translation.py", work)
    plan = read_json(work / "translation" / "plan.json")
    translations = {
        "title-node": "冻结讲义",
        "section-node": "1 基础",
        "theorem-node": "定理一。这是冻结的目标陈述。",
        "definition-node": "定义二。这是只包含锚点的目标陈述。",
        "paragraph-node": "本段结束这个合成讲义夹具。",
    }
    for batch in plan["batches"]:
        requests = read_jsonl(work / "translation" / batch["request_file"])
        responses = []
        for request in requests:
            placeholders = " ".join(
                token["placeholder"] for token in request.get("protected_tokens", [])
            )
            translation = translations[request["id"]]
            if placeholders:
                translation = f"{translation} {placeholders}"
            responses.append(
                {
                    "id": request["id"],
                    "source_sha256": request["source_sha256"],
                    "translation": translation,
                }
            )
        write_jsonl(work / "translation" / batch["response_file"], responses)
    run_work_script("audit_translation.py", work)
    run_work_script("build_outputs.py", work, "--basename", "fixture")
    run_work_script("audit_outputs.py", work)
    return ir["inventories"]["role_inventory"], work / "output" / "fixture.md"


def test_frozen_audit() -> None:
    profile = load_profile("lecture-notes-en-zh")
    with tempfile.TemporaryDirectory(prefix="v23-docx-test-") as temporary:
        work = Path(temporary)
        profile_path = work / "profile.json"
        write_json(profile_path, profile)
        theorem_source = "Theorem 1. Frozen source statement."
        definition_source = "Definition 2. Scoped anchor source statement."
        equation_source = "E equals m c squared."
        inventory, markdown_path = prepare_frozen_v2_output(
            work,
            profile,
            theorem_source,
            definition_source,
            equation_source,
        )
        output = work / "output"
        ir_path = work / "document-ir.json"

        docx_path = output / "fixture.docx"
        if shutil.which("pandoc"):
            build_process = subprocess.run(
                [
                    sys.executable,
                    str(SCRIPT_DIR / "build_docx.py"),
                    str(markdown_path),
                    str(docx_path),
                    "--work-dir",
                    str(work),
                    "--expected-role",
                    "theorem=1",
                    "--expected-role",
                    "definition=1",
                    "--expected-role",
                    "equation=1",
                    "--expected-problems",
                    "0",
                ],
                check=False,
                capture_output=True,
                text=True,
            )
            assert build_process.returncode == 0, build_process.stdout + build_process.stderr
        else:
            document = Document()
            document.styles.add_style("Block Text", WD_STYLE_TYPE.PARAGRAPH)
            add_generic_range(
                document,
                "theorem",
                "theorem",
                "theorem:theorem-node",
                theorem_source,
                "定理一。这是冻结的目标陈述。",
            )
            add_generic_range(
                document,
                "definition",
                "definition",
                "definition:definition-node",
                definition_source,
                "定义二。这是只包含锚点的目标陈述。",
                membership="anchor-only",
            )
            document.add_paragraph(equation_source)
            docx_style.configure_profile(profile)
            docx_style.apply_styles(
                document,
                document_title="V2.3 frozen audit",
                header_label="test",
                footer_label="测试",
            )
            document.save(docx_path)
        audit_path = output / "docx-audit.json"
        process = subprocess.run(
            [
                sys.executable,
                str(SCRIPT_DIR / "audit_docx.py"),
                str(docx_path),
                "--work-dir",
                str(work),
                "--expected-role",
                "theorem=1",
                "--expected-role",
                "definition=1",
                "--expected-role",
                "equation=1",
                "--output",
                str(audit_path),
            ],
            check=False,
            capture_output=True,
            text=True,
        )
        assert process.returncode == 0, process.stdout + process.stderr
        report = json.loads(audit_path.read_text(encoding="utf-8"))
        assert report["status"] == "passed"
        assert report["role_counts"]["theorem"] == 1
        assert report["role_counts"]["definition"] == 1
        assert report["scoped_anchor_callout_checks"] == {
            "definition:definition-node": True
        }
        assert report["non_structural_anchor_unboxed"][
            "equation:equation-node"
        ] is True
        assert report["source_only_occurrence_counts"] == {"equation-node": 1}
        assert report["document_ir_sha256"] == sha256_file(ir_path)
        passed_docx_bytes = docx_path.read_bytes()
        passed_audit_bytes = audit_path.read_bytes()

        def audit_fixture(path: Path) -> subprocess.CompletedProcess[str]:
            candidate_bytes = path.read_bytes()
            path.unlink()
            docx_path.write_bytes(candidate_bytes)
            try:
                return subprocess.run(
                    [
                        sys.executable,
                        str(SCRIPT_DIR / "audit_docx.py"),
                        str(docx_path),
                        "--work-dir",
                        str(work),
                        "--expected-role",
                        "theorem=1",
                        "--expected-role",
                        "definition=1",
                        "--expected-role",
                        "equation=1",
                    ],
                    check=False,
                    capture_output=True,
                    text=True,
                )
            finally:
                docx_path.write_bytes(passed_docx_bytes)
                audit_path.write_bytes(passed_audit_bytes)

        missing_document = Document()
        missing_document.styles.add_style("Block Text", WD_STYLE_TYPE.PARAGRAPH)
        add_generic_range(
            missing_document,
            "theorem",
            "theorem",
            "theorem:theorem-node",
            theorem_source,
            "定理一。这是冻结的目标陈述。",
        )
        missing_document.add_paragraph(definition_source)
        missing_document.add_paragraph("定义二。这是只包含锚点的目标陈述。")
        missing_document.add_paragraph(equation_source)
        docx_style.configure_profile(profile)
        docx_style.apply_styles(
            missing_document,
            document_title="missing scoped anchor",
            header_label="test",
            footer_label="测试",
        )
        missing_path = output / "missing-anchor-callout.docx"
        missing_document.save(missing_path)
        missing_audit = audit_fixture(missing_path)
        assert missing_audit.returncode == 1
        assert "scoped_anchor_callouts_are_structurally_stable" in missing_audit.stdout

        absorbed_document = Document()
        absorbed_document.styles.add_style("Block Text", WD_STYLE_TYPE.PARAGRAPH)
        add_generic_range(
            absorbed_document,
            "theorem",
            "theorem",
            "theorem:theorem-node",
            theorem_source,
            "定理一。这是冻结的目标陈述。",
        )
        absorbed_document.add_paragraph(
            f"{GENERIC_BEGIN}|definition|definition|anchor-only|definition:definition-node",
            style="Block Text",
        )
        absorbed_document.add_paragraph(definition_source, style="Block Text")
        absorbed_document.add_paragraph(equation_source, style="Block Text")
        absorbed_document.add_paragraph("", style="Block Text")
        absorbed_document.add_paragraph(
            "定义二。这是只包含锚点的目标陈述。", style="Block Text"
        )
        absorbed_document.add_paragraph(
            f"{GENERIC_END}|definition|definition|anchor-only|definition:definition-node",
            style="Block Text",
        )
        docx_style.configure_profile(profile)
        docx_style.apply_styles(
            absorbed_document,
            document_title="absorbed neighbor",
            header_label="test",
            footer_label="测试",
        )
        absorbed_path = output / "absorbed-neighbor.docx"
        absorbed_document.save(absorbed_path)
        absorbed_audit = audit_fixture(absorbed_path)
        assert absorbed_audit.returncode == 1
        assert "scoped_anchor_callouts_are_structurally_stable" in absorbed_audit.stdout
        assert "non_structural_anchor_groups_are_not_boxed" in absorbed_audit.stdout

        compile_context = compile_docx_pdf.load_v2_context(work, docx_path)
        assert compile_context["schema_version"] == 2
        assert compile_context["build"]["role_inventory"] == inventory
        assert compile_context["docx_audit"]["status"] == "passed"
        assert compile_context["docx_audit_bindings"]["docx_sha256"] == sha256_file(
            docx_path
        )
        assert compile_docx_pdf.parse_expected_roles(["theorem=1", "equation=1"]) == {
            "theorem": 1,
            "equation": 1,
        }

        audit_bytes = audit_path.read_bytes()
        audit_path.unlink()
        try:
            compile_docx_pdf.load_v2_context(work, docx_path)
        except ValueError as exc:
            assert "missing DOCX audit" in str(exc)
        else:
            raise AssertionError("schema V2 compile accepted a missing DOCX audit")
        audit_path.write_bytes(audit_bytes)

        docx_bytes = docx_path.read_bytes()
        docx_path.write_bytes(docx_bytes + b"changed after DOCX audit")
        try:
            compile_docx_pdf.load_v2_context(work, docx_path)
        except ValueError as exc:
            assert "docx_sha256" in str(exc)
        else:
            raise AssertionError("schema V2 compile accepted modified DOCX bytes")
        docx_path.write_bytes(docx_bytes)

        stale_audit = json.loads(audit_bytes.decode("utf-8"))
        stale_audit["profile_file_sha256"] = "0" * 64
        write_json(audit_path, stale_audit)
        try:
            compile_docx_pdf.load_v2_context(work, docx_path)
        except ValueError as exc:
            assert "profile_file_sha256" in str(exc)
        else:
            raise AssertionError("schema V2 compile accepted stale audit bindings")
        audit_path.write_bytes(audit_bytes)


def main() -> None:
    test_legacy_transform()
    test_generic_ast()
    test_shared_style_roles()
    test_html_table_materialization()
    test_frozen_audit()
    print(
        "V2.3 DOCX tests passed: legacy, structural AST, shared styles, "
        "native tables, frozen audit"
    )


if __name__ == "__main__":
    main()
